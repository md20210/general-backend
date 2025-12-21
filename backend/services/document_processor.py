"""Document Processing Service for PDFs, DOCX, and URLs."""
from typing import Dict, BinaryIO
import io

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("WARNING: PyPDF2 not installed - PDF processing disabled")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not installed - DOCX processing disabled")

try:
    import requests
    from bs4 import BeautifulSoup
    WEB_SCRAPING_AVAILABLE = True
except ImportError:
    WEB_SCRAPING_AVAILABLE = False
    print("WARNING: requests/beautifulsoup4 not installed - URL scraping disabled")


class DocumentProcessor:
    """Process documents from various sources (PDF, DOCX, URL, text)."""

    @staticmethod
    def extract_from_pdf(file: BinaryIO) -> str:
        """
        Extract text from PDF file.

        Args:
            file: Binary file object (from file upload)

        Returns:
            Extracted text content
        """
        if not PYPDF2_AVAILABLE:
            raise RuntimeError("PyPDF2 not installed")

        text = ""
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + "\n"

        return text.strip()

    @staticmethod
    def extract_from_docx(file: BinaryIO) -> str:
        """
        Extract text from DOCX file.

        Args:
            file: Binary file object (from file upload)

        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")

        doc = DocxDocument(file)
        text = ""

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text.strip()

    @staticmethod
    def scrape_website(url: str, max_length: int = 10000) -> Dict[str, str]:
        """
        Scrape website and extract content.

        Args:
            url: Website URL
            max_length: Maximum content length (default: 10000 chars)

        Returns:
            Dict with title, description, content
        """
        if not WEB_SCRAPING_AVAILABLE:
            raise RuntimeError("requests/beautifulsoup4 not installed")

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')

        # Remove script and style tags
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        # Extract title
        title = soup.title.string if soup.title else ""

        # Extract meta description
        description = ""
        meta_desc = soup.find("meta", attrs={"name": "description"})
        if meta_desc:
            description = meta_desc.get("content", "")

        # Extract main content
        content = soup.get_text(separator="\n", strip=True)

        return {
            "title": title.strip(),
            "description": description.strip(),
            "content": content[:max_length].strip()
        }

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: Text to chunk
            chunk_size: Size of each chunk
            overlap: Overlap between chunks

        Returns:
            List of text chunks
        """
        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap

        return chunks


# Global instance
document_processor = DocumentProcessor()
