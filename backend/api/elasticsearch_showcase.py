"""Elasticsearch Showcase API Router."""
from datetime import datetime
from typing import List, Optional
from uuid import UUID
import fastapi
from fastapi import APIRouter, Depends, HTTPException, Query, File, UploadFile
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession
from backend.database import get_async_session
from backend.auth.dependencies import current_active_user
from backend.models.user import User
from backend.models.elasticsearch_showcase import UserElasticProfile, ElasticJobAnalysis
from backend.schemas.elasticsearch_showcase import (
    UserProfileRequest,
    UserProfileResponse,
    JobAnalysisRequest,
    JobAnalysisResponse,
    ComparisonReport,
    AdvancedFeatures,
    StatusResponse,
    ErrorResponse,
)
from backend.services.elasticsearch_service import ElasticsearchService
from backend.services.elasticsearch_comparison_service import ElasticsearchComparisonService
from backend.services.llm_gateway import LLMGateway
from backend.services.logstash_service import LogstashService
from backend.services.demo_data_generator import DemoDataGenerator
from backend.services.elasticsearch_vector_service import ElasticsearchVectorService
import logging
import json
import re
import os
import httpx

logger = logging.getLogger(__name__)

router = APIRouter(
    prefix="/elasticsearch",
    tags=["Elasticsearch Showcase"],
)

# Initialize services
es_service = ElasticsearchService()
comparison_service = ElasticsearchComparisonService()
llm_gateway = LLMGateway()
logstash_service = LogstashService(logstash_url=os.getenv("LOGSTASH_URL"))
demo_generator = DemoDataGenerator()
vector_service = ElasticsearchVectorService()  # pgvector instead of ChromaDB!


# ============================================================================
# Helper Functions
# ============================================================================

async def crawl_url(url: str, max_length: int = 10000) -> str:
    """
    Crawl URL and extract clean text content.

    Args:
        url: The URL to crawl
        max_length: Maximum length of extracted text (default: 10000 chars)

    Returns:
        Cleaned text content from the URL, or empty string if crawling fails
    """
    try:
        # Add https:// if URL doesn't have a protocol
        if not url.startswith(('http://', 'https://')):
            url = f'https://{url}'

        async with httpx.AsyncClient(
            timeout=30.0,
            follow_redirects=True,
            headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
        ) as client:
            response = await client.get(url)

            if response.status_code != 200:
                logger.warning(f"URL {url} returned status code {response.status_code}")
                return ""

            # Try to parse HTML
            try:
                from bs4 import BeautifulSoup

                # Parse HTML
                soup = BeautifulSoup(response.text, 'html.parser')

                # Remove script, style, nav, footer, header elements
                for element in soup(["script", "style", "nav", "footer", "header", "aside", "iframe", "noscript"]):
                    element.decompose()

                # Get text from main content areas (prioritize main, article, section)
                main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|main|body'))

                if main_content:
                    text = main_content.get_text(separator=' ', strip=True)
                else:
                    text = soup.get_text(separator=' ', strip=True)

                # Clean whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)

                # Truncate if too long
                if len(text) > max_length:
                    text = text[:max_length] + "..."

                return text

            except ImportError:
                # BeautifulSoup not available, return raw text (fallback)
                logger.warning("BeautifulSoup not available, using raw text extraction")
                text = response.text[:max_length]
                # Simple cleanup: remove HTML tags
                text = re.sub(r'<[^>]+>', ' ', text)
                text = re.sub(r'\s+', ' ', text).strip()
                return text

    except httpx.TimeoutException:
        logger.error(f"Timeout while crawling {url}")
        return ""
    except Exception as e:
        logger.error(f"Failed to crawl {url}: {e}")
        return ""


def deduplicate_text_content(text: str, similarity_threshold: float = 0.75) -> str:
    """
    Remove semantically similar/duplicate sentences from text.

    This prevents redundant information from being indexed multiple times
    when the same fact appears in CV, Homepage, and LinkedIn profile.

    Args:
        text: Input text that may contain duplicates
        similarity_threshold: Jaccard similarity threshold (0-1) for considering sentences similar

    Returns:
        Deduplicated text with unique sentences only
    """
    # Split into sentences
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]

    if not sentences:
        return text

    # Track unique sentences
    unique_sentences = []
    seen_words_sets = []

    for sentence in sentences:
        # Convert to word set (lowercase, alphanumeric only)
        words = set(re.findall(r'\b\w+\b', sentence.lower()))

        if not words:
            continue

        # Check similarity with already seen sentences
        is_duplicate = False
        for seen_words in seen_words_sets:
            # Jaccard similarity = intersection / union
            intersection = len(words & seen_words)
            union = len(words | seen_words)
            similarity = intersection / union if union > 0 else 0

            if similarity >= similarity_threshold:
                is_duplicate = True
                break

        if not is_duplicate:
            unique_sentences.append(sentence)
            seen_words_sets.append(words)

    # Reconstruct text
    return '. '.join(unique_sentences) + '.'


def extract_skills_from_cv(cv_text: str) -> List[str]:
    """Extract skills from CV text using simple keyword matching."""
    # Common skills to look for (case-insensitive)
    common_skills = [
        "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Rust",
        "React", "Vue", "Angular", "Node.js", "Django", "Flask", "FastAPI",
        "Docker", "Kubernetes", "AWS", "Azure", "GCP",
        "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis",
        "Git", "CI/CD", "Jenkins", "GitHub Actions",
        "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch",
        "REST API", "GraphQL", "Microservices",
        "Agile", "Scrum", "DevOps",
        "Linux", "Bash", "Shell"
    ]

    found_skills = []
    cv_lower = cv_text.lower()

    for skill in common_skills:
        if skill.lower() in cv_lower:
            found_skills.append(skill)

    return found_skills


def extract_experience_years(cv_text: str) -> Optional[int]:
    """Extract years of experience from CV text."""
    # Look for patterns like "X years", "X+ years", "X-Y years"
    patterns = [
        r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
        r'experience\s*:\s*(\d+)\+?\s*years?',
        r'(\d+)\+?\s*years?\s+in\s+(?:the\s+)?industry'
    ]

    for pattern in patterns:
        match = re.search(pattern, cv_text.lower())
        if match:
            return int(match.group(1))

    return None


def extract_education_level(cv_text: str) -> Optional[str]:
    """Extract education level from CV text."""
    cv_lower = cv_text.lower()

    if "phd" in cv_lower or "ph.d" in cv_lower or "doctorate" in cv_lower:
        return "PhD"
    elif "master" in cv_lower or "msc" in cv_lower or "m.sc" in cv_lower:
        return "Master"
    elif "bachelor" in cv_lower or "bsc" in cv_lower or "b.sc" in cv_lower:
        return "Bachelor"
    elif "associate" in cv_lower:
        return "Associate"

    return None


def extract_job_titles(cv_text: str) -> List[str]:
    """Extract job titles from CV text."""
    # Common job title patterns
    title_keywords = [
        "Software Engineer", "Developer", "Programmer",
        "Senior Engineer", "Lead Engineer", "Principal Engineer",
        "Data Scientist", "Data Engineer", "ML Engineer",
        "DevOps Engineer", "Site Reliability Engineer",
        "Full Stack Developer", "Frontend Developer", "Backend Developer",
        "Engineering Manager", "Technical Lead", "Team Lead",
        "Architect", "Solutions Architect", "System Architect"
    ]

    found_titles = []
    cv_lower = cv_text.lower()

    for title in title_keywords:
        if title.lower() in cv_lower:
            found_titles.append(title)

    return found_titles[:5]  # Return max 5 titles


async def extract_cv_info_with_llm(cv_text: str, provider: str = "grok") -> dict:
    """
    Extract ALL CV information using LLM in a single call for better accuracy and efficiency.

    Returns dict with: skills, experience_years, education_level, job_titles
    Falls back to regex-based extraction if LLM fails.
    """
    prompt = f"""Analyze this CV and extract the following information. Return ONLY valid JSON, no explanations.

CV:
{cv_text[:4000]}

Extract:
1. skills: Array of all technical and professional skills (max 50)
2. experience_years: Total years of professional experience (integer or null)
3. education_level: Highest education ("PhD", "Master", "Bachelor", "Associate", or null)
4. job_titles: Array of job titles mentioned in CV (max 5, most recent/relevant first)

Return EXACTLY this JSON format:
{{
  "skills": ["skill1", "skill2", ...],
  "experience_years": 10,
  "education_level": "Master",
  "job_titles": ["Senior Engineer", "Developer", ...]
}}
"""

    try:
        logger.info(f"Extracting CV info with LLM ({provider})...")
        response_dict = llm_gateway.generate(prompt=prompt, provider=provider, temperature=0.3)
        response = response_dict.get("response", "")

        # Extract JSON from response
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            cv_info = json.loads(json_match.group(0))

            # Validate and clean data
            result = {
                "skills": cv_info.get("skills", [])[:50],
                "experience_years": cv_info.get("experience_years"),
                "education_level": cv_info.get("education_level"),
                "job_titles": cv_info.get("job_titles", [])[:5]
            }

            logger.info(f"✅ LLM extracted: {len(result['skills'])} skills, "
                       f"{result['experience_years']} years exp, "
                       f"{result['education_level']} education")
            return result
        else:
            logger.warning("LLM didn't return valid JSON, using fallback")
            raise ValueError("Invalid JSON response")
    except Exception as e:
        logger.error(f"Error extracting CV info with LLM: {e}")
        logger.info("Falling back to regex-based extraction")
        return {
            "skills": extract_skills_from_cv(cv_text),
            "experience_years": extract_experience_years(cv_text),
            "education_level": extract_education_level(cv_text),
            "job_titles": extract_job_titles(cv_text)
        }


async def extract_skills_with_llm(cv_text: str, provider: str = "grok") -> List[str]:
    """
    Extract skills from CV using LLM (deprecated - use extract_cv_info_with_llm instead).
    Kept for backward compatibility.
    """
    cv_info = await extract_cv_info_with_llm(cv_text, provider)
    return cv_info.get("skills", [])


async def analyze_job_with_llm(job_description: str, cv_text: str, provider: str = "grok") -> dict:
    """Analyze job description and compare with CV using LLM."""
    prompt = f"""Analyze this job description and compare it with the candidate's CV.
BE REALISTIC and OBJECTIVE. Do NOT exaggerate qualifications. Base analysis ONLY on factual evidence in the CV.

JOB DESCRIPTION:
{job_description}

CANDIDATE CV:
{cv_text}

CRITICAL INSTRUCTIONS:
1. Extract ONLY information explicitly stated in the documents
2. For skills: Categorize as Exact/Strong/Partial/Missing match
3. For experience: Compare requested years vs candidate's actual years
4. Be honest about overqualification or underqualification
5. NO marketing language or exaggeration

Return ONLY valid JSON in this exact format:
{{
  "job_analysis": {{
    "company": "Company name from job posting",
    "role": "Exact job title from posting",
    "location": "Location or Remote",
    "remote_policy": "Remote/Hybrid/On-site",
    "seniority": "Junior/Mid/Senior/Lead/Principal",
    "salary_range": {{"min": null, "max": null, "currency": "EUR"}},
    "requirements": {{
      "must_have": ["skill1", "skill2"],
      "nice_to_have": ["skill3"],
      "years_experience": {{"min": 5, "max": 10}},
      "education": "Bachelor/Master/PhD or not specified",
      "languages": ["German", "English"],
      "certifications": []
    }},
    "responsibilities": ["responsibility1", "responsibility2"],
    "keywords": ["keyword1", "keyword2"],
    "red_flags": ["Concerns like overqualification, salary mismatch, relocation"],
    "green_flags": ["Genuine strengths from CV"]
  }},
  "fit_score": {{
    "total": 85,
    "breakdown": {{
      "experience_match": 90,
      "skills_match": 85,
      "education_match": 100,
      "location_match": 80,
      "salary_match": 75,
      "culture_match": 80,
      "role_type_match": 90
    }},
    "component_explanations": {{
      "experience_match": "Requested: 5-7 years in Search AI | Candidate: 25+ years (significantly overqualified)",
      "skills_match": "Has 6/8 core skills. Missing: Kubernetes, Terraform",
      "education_match": "Requested: Bachelor | Candidate: Master KIT (exceeds requirement)",
      "location_match": "Job: Barcelona | Candidate: Barcelona (perfect match)",
      "salary_match": "Candidate likely expects senior compensation above role budget",
      "culture_match": "Enterprise background aligns with company needs",
      "role_type_match": "Architect role matches candidate seniority"
    }},
    "exact_match_skills": ["Elasticsearch", "Logstash", "Kibana", "Python", "FastAPI"],
    "strong_match_skills": ["RAG (Job mentions Search AI)", "Vector DBs (Job mentions Semantic Search)"],
    "partial_match_skills": ["TOGAF (Job mentions Architecture)", "Scrum (Job mentions Agile)"],
    "missing_skills": ["Kubernetes", "Terraform"]
  }},
  "interview_success": {{
    "probability": 92,
    "interpretation": "Interview Success Probability - likelihood of getting past screening based on resume match",
    "factors": [
      {{"factor": "Native German speaker (explicit requirement)", "impact": 20}},
      {{"factor": "Elastic Stack hands-on experience (showcase demo)", "impact": 25}},
      {{"factor": "25 years enterprise experience (strong but possibly overqualified)", "impact": 15}},
      {{"factor": "Master degree from KIT (exceeds Bachelor requirement)", "impact": 10}},
      {{"factor": "May be too senior for role level and salary", "impact": -15}},
      {{"factor": "Exact location match (Barcelona)", "impact": 10}},
      {{"factor": "Proven customer success track record", "impact": 15}},
      {{"factor": "Missing DevOps tools (K8s, Terraform)", "impact": -8}}
    ],
    "recommendation": "High interview probability. Address overqualification concern in cover letter. Emphasize passion for customer-facing role."
  }}
}}
"""

    try:
        response_dict = llm_gateway.generate(prompt=prompt, provider=provider)
        response = response_dict.get("response", "")

        # Extract JSON from response
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            analysis = json.loads(json_match.group(0))
            return analysis
        else:
            logger.error("LLM didn't return valid JSON for job analysis")
            # Return default structure
            return {
                "job_analysis": {
                    "company": "Unknown",
                    "role": "Position",
                    "location": "Unknown",
                    "remote_policy": "Unknown",
                    "seniority": "Unknown",
                    "salary_range": {"min": None, "max": None, "currency": "EUR"},
                    "requirements": {
                        "must_have": [],
                        "nice_to_have": [],
                        "years_experience": {"min": 0, "max": None},
                        "education": "Not specified",
                        "languages": [],
                        "certifications": []
                    },
                    "responsibilities": [],
                    "keywords": [],
                    "red_flags": [],
                    "green_flags": []
                },
                "fit_score": {
                    "total": 0,
                    "breakdown": {
                        "experience_match": 0,
                        "skills_match": 0,
                        "education_match": 0,
                        "location_match": 0,
                        "salary_match": 0,
                        "culture_match": 0,
                        "role_type_match": 0
                    },
                    "component_explanations": {
                        "experience_match": "Unable to analyze",
                        "skills_match": "Unable to analyze",
                        "education_match": "Unable to analyze",
                        "location_match": "Unable to analyze",
                        "salary_match": "Unable to analyze",
                        "culture_match": "Unable to analyze",
                        "role_type_match": "Unable to analyze"
                    },
                    "exact_match_skills": [],
                    "strong_match_skills": [],
                    "partial_match_skills": [],
                    "missing_skills": []
                },
                "interview_success": {
                    "probability": 0,
                    "interpretation": "Interview Success Probability",
                    "factors": [],
                    "recommendation": "Unable to analyze - please try again"
                }
            }
    except Exception as e:
        logger.error(f"Error analyzing job with LLM: {e}")
        raise


# ============================================================================
# Profile Endpoints
# ============================================================================

@router.post("/profile", response_model=UserProfileResponse)
async def create_or_update_profile(
    profile_data: UserProfileRequest,
    provider: str = Query("grok", description="LLM provider for skill extraction (ollama, grok, anthropic)"),
    skip_processing: bool = Query(False, description="Skip LLM processing and text deduplication (URLs still crawled)"),
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """
    Create or update user profile with CV and documents.

    This endpoint:
    1. Stores CV, cover letter, and URLs in PostgreSQL
    2. Optionally extracts skills, experience, education using LLM (skip_processing=False)
    3. Crawls URLs for additional content (always performed)
    4. Optionally deduplicates text content (skip_processing=False)
    5. Indexes data in Elasticsearch for fast searching

    Args:
        skip_processing: If True, skips LLM analysis and text deduplication.
                        URLs are still crawled but content is not AI-processed.
    """
    try:
        # Store user_id early to avoid lazy loading after rollback
        user_id = str(user.id)

        # Initialize import status tracking
        elasticsearch_indexed = False
        pgvector_chunks = 0
        homepage_crawled = False
        linkedin_crawled = False

        logger.info(f"🚀 Starting CV import for user {user_id} (skip_processing={skip_processing})")

        # Conditionally extract CV info with LLM
        if skip_processing:
            logger.info(f"📋 Step 1/7: Skipping LLM extraction (raw import mode)")
            skills = []
            experience_years = None
            education_level = None
            job_titles = []
        else:
            logger.info(f"📋 Step 1/7: Extracting CV information with LLM provider: {provider}")
            # Extract ALL information from CV using LLM (more accurate than regex)
            cv_info = await extract_cv_info_with_llm(profile_data.cv_text, provider)
            skills = cv_info.get("skills", [])
            experience_years = cv_info.get("experience_years")
            education_level = cv_info.get("education_level")
            job_titles = cv_info.get("job_titles", [])

        logger.info(f"📋 Step 2/7: Saving profile to database...")
        # Check if profile exists
        statement = select(UserElasticProfile).where(UserElasticProfile.user_id == user_id)
        result = await db.execute(statement)
        existing_profile = result.scalars().first()

        if existing_profile:
            # Update existing profile
            existing_profile.cv_text = profile_data.cv_text
            existing_profile.cover_letter_text = profile_data.cover_letter_text
            existing_profile.homepage_url = profile_data.homepage_url
            existing_profile.linkedin_url = profile_data.linkedin_url
            existing_profile.skills_extracted = skills
            existing_profile.experience_years = experience_years
            existing_profile.education_level = education_level
            existing_profile.job_titles = job_titles
            existing_profile.updated_at = datetime.utcnow()

            db.add(existing_profile)
            await db.commit()
            await db.refresh(existing_profile)
            profile = existing_profile
        else:
            # Create new profile
            new_profile = UserElasticProfile(
                user_id=user_id,
                cv_text=profile_data.cv_text,
                cover_letter_text=profile_data.cover_letter_text,
                homepage_url=profile_data.homepage_url,
                linkedin_url=profile_data.linkedin_url,
                skills_extracted=skills,
                experience_years=experience_years,
                education_level=education_level,
                job_titles=job_titles,
            )
            db.add(new_profile)
            await db.commit()
            await db.refresh(new_profile)
            profile = new_profile

        logger.info(f"📋 Step 3/7: Deleting old vector data...")
        # Delete existing vector data to avoid duplicates (pgvector only)
        # Note: If this fails, we skip all pgvector operations to avoid session errors
        pgvector_available = False
        logger.info(f"Deleting existing vector data for user {user_id}")
        try:
            # Delete from pgvector (ChromaDB replacement)
            if vector_service.is_available():
                await vector_service.delete_collection(session=db, user_id=UUID(user_id), project_id=None)
                logger.info(f"✅ Deleted existing pgvector data for user {user_id}")
                pgvector_available = True  # Only set to True if delete succeeded
        except Exception as e:
            logger.warning(f"⚠️  Failed to delete pgvector data: {e}")
            logger.warning(f"⚠️  Skipping all pgvector operations (session may be in error state)")
            # Don't manipulate session - it causes greenlet_spawn errors

        try:
            # Delete from Elasticsearch
            await es_service.delete_user_cv_data(user_id=user_id)
            logger.info(f"✅ Deleted existing Elasticsearch data for user {user_id}")
        except Exception as e:
            logger.warning(f"⚠️  Failed to delete Elasticsearch data: {e}")

        logger.info(f"📋 Step 4/7: Crawling URLs (Homepage + LinkedIn)...")
        # Crawl URLs and extract additional content
        additional_content = ""

        # Crawl Homepage URL
        if profile_data.homepage_url:
            try:
                logger.info(f"Crawling homepage: {profile_data.homepage_url}")
                homepage_content = await crawl_url(profile_data.homepage_url)
                if homepage_content:
                    additional_content += f"\n\n=== Homepage Content ({profile_data.homepage_url}) ===\n{homepage_content}\n"
                    logger.info(f"✅ Successfully crawled homepage: {len(homepage_content)} chars")
                    homepage_crawled = True
            except Exception as e:
                logger.warning(f"⚠️  Failed to crawl homepage: {e}")

        # Crawl LinkedIn URL
        if profile_data.linkedin_url:
            try:
                logger.info(f"Crawling LinkedIn: {profile_data.linkedin_url}")
                linkedin_content = await crawl_url(profile_data.linkedin_url)
                if linkedin_content:
                    additional_content += f"\n\n=== LinkedIn Profile ({profile_data.linkedin_url}) ===\n{linkedin_content}\n"
                    logger.info(f"✅ Successfully crawled LinkedIn: {len(linkedin_content)} chars")
                    linkedin_crawled = True
            except Exception as e:
                logger.warning(f"⚠️  Failed to crawl LinkedIn: {e}")

        # Combine all content for indexing
        full_cv_content = profile_data.cv_text
        if profile_data.cover_letter_text:
            full_cv_content += f"\n\n=== Cover Letter ===\n{profile_data.cover_letter_text}"
        if additional_content:
            full_cv_content += additional_content

        # Deduplicate content to avoid redundant information (only if not skipping processing)
        # (e.g., "Cognizant" appears in CV, Homepage, AND LinkedIn)
        if skip_processing:
            logger.info(f"Skipping deduplication (raw import mode) - content length: {len(full_cv_content)} chars")
            full_cv_content_deduplicated = full_cv_content
        else:
            logger.info(f"Original content length: {len(full_cv_content)} chars")
            full_cv_content_deduplicated = deduplicate_text_content(full_cv_content, similarity_threshold=0.75)
            logger.info(f"Deduplicated content length: {len(full_cv_content_deduplicated)} chars")
            logger.info(f"Removed {len(full_cv_content) - len(full_cv_content_deduplicated)} chars of duplicate content")

        logger.info(f"📋 Step 5/7: Indexing in Elasticsearch...")
        # Index in Elasticsearch FIRST (before profile update)
        # This way, if profile update fails due to aborted transaction, Elasticsearch is already indexed
        await es_service.index_cv_data(
            user_id=user_id,
            cv_text=full_cv_content_deduplicated,  # Deduplicated content (no redundant sentences)
            skills=skills,
            experience_years=experience_years,
            education_level=education_level,
            job_titles=job_titles,
            homepage_url=profile_data.homepage_url,
            linkedin_url=profile_data.linkedin_url
        )
        logger.info(f"✅ Indexed CV in Elasticsearch for user {user_id}")
        elasticsearch_indexed = True

        logger.info(f"📋 Step 6/7: Updating profile with crawled content...")
        # Update profile with full content (including crawled data)
        # Note: If pgvector delete failed, transaction may be in aborted state
        # In that case, skip the profile update (Elasticsearch is already indexed)
        # (Profile already has basic CV data from initial creation)
        session_aborted = False
        try:
            profile.cv_text = full_cv_content_deduplicated
            db.add(profile)
            await db.commit()
            await db.refresh(profile)
            logger.info(f"✅ Updated profile cv_text with crawled content: {len(full_cv_content_deduplicated)} chars")
        except Exception as e:
            logger.warning(f"⚠️ Profile cv_text update failed (transaction aborted): {e}")
            logger.warning(f"⚠️ Skipping profile update (Elasticsearch already indexed)")
            # Rollback to clear aborted transaction state
            await db.rollback()
            session_aborted = True
            # No more awaits after rollback - just continue to logging and return

        logger.info(f"📋 Step 7/7: Indexing in pgvector...")
        # Index in pgvector (for fair comparison with Elasticsearch)
        # Only if delete succeeded AND session not aborted (otherwise can't use db session)
        if pgvector_available and not session_aborted:
            try:
                if vector_service.is_available():
                    logger.info(f"📝 Preparing pgvector indexing for user {user_id}...")
                    # Prepare document content with all CV information (deduplicated!)
                    cv_content = f"""
{full_cv_content_deduplicated}

Skills: {', '.join(skills) if skills else 'N/A'}
Experience: {experience_years if experience_years else 'N/A'} years
Education: {education_level if education_level else 'N/A'}
Job Titles: {', '.join(job_titles) if job_titles else 'N/A'}
                    """.strip()

                    # Add to pgvector with metadata
                    metadata = {
                        "type": "cv",
                        "skills": ",".join(skills) if skills else "",
                        "job_titles": ",".join(job_titles) if job_titles else "",
                    }
                    # Only add non-None values
                    if experience_years is not None:
                        metadata["experience_years"] = str(experience_years)
                    if education_level is not None:
                        metadata["education_level"] = education_level
                    if profile_data.homepage_url:
                        metadata["homepage_url"] = profile_data.homepage_url
                    if profile_data.linkedin_url:
                        metadata["linkedin_url"] = profile_data.linkedin_url

                    logger.info(f"📦 Adding documents to pgvector (chunk_size=500)...")
                    pgvector_chunks = await vector_service.add_documents(
                        session=db,
                        user_id=UUID(user_id),
                        documents=[{
                            "id": f"cv_{user_id}",
                            "content": cv_content,
                            "metadata": metadata
                        }],
                        project_id=None,  # Use global collection for elasticsearch showcase
                        chunk_size=500
                    )
                    logger.info(f"✅ Indexed CV in pgvector: {pgvector_chunks} chunks added for user {user_id}")
                else:
                    logger.warning("⚠️  pgvector not available - skipping vector indexing")
            except Exception as e:
                logger.error(f"❌ pgvector add_documents failed (continuing anyway): {e}")
                # Profile is already committed - just log the error and continue
                # Don't fail the request if pgvector fails - Elasticsearch indexing succeeded
        else:
            if session_aborted:
                logger.warning("⚠️  Skipping pgvector add_documents (session aborted after transaction error)")
            else:
                logger.warning("⚠️  Skipping pgvector add_documents (delete failed or unavailable)")

        # Log success summary with important metrics
        logger.info(f"✅ Profile import completed for user {user_id}")
        logger.info(f"📊 Import Summary:")
        logger.info(f"  - Skills extracted: {len(skills)}")
        logger.info(f"  - Experience: {experience_years} years" if experience_years else "  - Experience: N/A")
        logger.info(f"  - Education: {education_level}" if education_level else "  - Education: N/A")
        logger.info(f"  - Job titles: {len(job_titles)}")
        logger.info(f"  - CV content: {len(full_cv_content_deduplicated)} chars (deduplicated)")
        logger.info(f"  - Elasticsearch: ✅ Indexed successfully")
        if pgvector_chunks > 0:
            logger.info(f"  - pgvector: ✅ Indexed {pgvector_chunks} chunks")
        else:
            logger.info(f"  - pgvector: ⚠️ Skipped (session error or unavailable)")
        logger.info(f"🎉 Import process completed successfully!")

        # Re-query the profile to ensure it's attached to a valid session
        # This prevents greenlet_spawn errors when FastAPI serializes the response
        fresh_statement = select(UserElasticProfile).where(UserElasticProfile.user_id == user_id)
        fresh_result = await db.execute(fresh_statement)
        fresh_profile = fresh_result.scalars().first()

        if not fresh_profile:
            logger.error(f"❌ Failed to retrieve profile after import for user {user_id}")
            raise HTTPException(status_code=500, detail="Profile import succeeded but failed to retrieve result")

        # Add import status information to response
        fresh_profile.elasticsearch_indexed = elasticsearch_indexed
        fresh_profile.pgvector_chunks = pgvector_chunks
        fresh_profile.homepage_crawled = homepage_crawled
        fresh_profile.linkedin_crawled = linkedin_crawled

        return fresh_profile

    except Exception as e:
        logger.error(f"❌ Error creating/updating profile for user: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {str(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/profile", response_model=UserProfileResponse)
async def get_profile(
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """Get user profile."""
    statement = select(UserElasticProfile).where(UserElasticProfile.user_id == str(user.id))
    result = await db.execute(statement)
    profile = result.scalars().first()

    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")

    return profile


# ============================================================================
# Job Analysis Endpoints
# ============================================================================

@router.post("/analyze", response_model=JobAnalysisResponse)
async def analyze_job(
    analysis_request: JobAnalysisRequest,
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """
    Analyze job posting and compare ChromaDB vs Elasticsearch.

    This endpoint:
    1. Runs search in both ChromaDB and Elasticsearch
    2. Compares performance metrics
    3. Demonstrates advanced Elasticsearch features
    4. Stores results in database
    """
    try:
        # Check if profile exists
        statement = select(UserElasticProfile).where(UserElasticProfile.user_id == str(user.id))
        result = await db.execute(statement)
        profile = result.scalars().first()

        if not profile:
            raise HTTPException(
                status_code=404,
                detail="Profile not found. Please create a profile first by uploading your CV."
            )

        # Get job description - either from request or by crawling URL
        job_description = analysis_request.job_description
        if not job_description and analysis_request.job_url:
            # Crawl job URL to get description
            logger.info(f"Crawling job URL: {analysis_request.job_url}")
            try:
                crawled_text = await crawl_url(analysis_request.job_url)
                if crawled_text:
                    job_description = crawled_text
                    logger.info(f"Successfully crawled job description: {len(job_description)} chars")
                else:
                    logger.warning(f"URL crawling returned empty content")
            except Exception as e:
                logger.error(f"Failed to crawl job URL: {e}")

        if not job_description:
            raise HTTPException(
                status_code=400,
                detail="Either job_description or job_url must be provided"
            )

        # Perform LLM-based job analysis (with fallback to default values)
        logger.info(f"Starting LLM job analysis for user {user.id}")
        try:
            llm_analysis = await analyze_job_with_llm(
                job_description=job_description,
                cv_text=profile.cv_text,
                provider=analysis_request.provider or "grok"
            )
            logger.info(f"LLM job analysis completed for user {user.id}")
        except Exception as e:
            logger.error(f"LLM analysis failed, using default values: {e}")
            # Use default empty analysis if LLM fails
            llm_analysis = {
                "job_analysis": {
                    "company": "Unknown",
                    "role": "Position",
                    "location": "Unknown",
                    "remote_policy": "Unknown",
                    "seniority": "Unknown",
                    "salary_range": {"min": None, "max": None, "currency": "EUR"},
                    "requirements": {
                        "must_have": [],
                        "nice_to_have": [],
                        "years_experience": {"min": 0, "max": None},
                        "education": "Not specified",
                        "languages": [],
                        "certifications": []
                    },
                    "responsibilities": [],
                    "keywords": [],
                    "red_flags": [],
                    "green_flags": []
                },
                "fit_score": {
                    "total": 0,
                    "breakdown": {
                        "experience_match": 0,
                        "skills_match": 0,
                        "education_match": 0,
                        "location_match": 0,
                        "salary_match": 0,
                        "culture_match": 0,
                        "role_type_match": 0
                    },
                    "component_explanations": {
                        "experience_match": "Unable to analyze",
                        "skills_match": "Unable to analyze",
                        "education_match": "Unable to analyze",
                        "location_match": "Unable to analyze",
                        "salary_match": "Unable to analyze",
                        "culture_match": "Unable to analyze",
                        "role_type_match": "Unable to analyze"
                    },
                    "exact_match_skills": [],
                    "strong_match_skills": [],
                    "partial_match_skills": [],
                    "missing_skills": []
                },
                "interview_success": {
                    "probability": 0,
                    "interpretation": "Interview Success Probability",
                    "factors": [],
                    "recommendation": "Unable to analyze - LLM service unavailable"
                }
            }

        # Run comparison (with fallback to empty results)
        try:
            comparison_results = await comparison_service.compare_search_performance(
                user_id=str(user.id),
                job_description=job_description,
                required_skills=analysis_request.required_skills or []
            )
        except Exception as e:
            logger.warning(f"Comparison service failed, using empty results: {e}")
            comparison_results = {
                "chromadb": {"results": {}, "search_time_ms": 0, "total_matches": 0, "relevance_scores": []},
                "elasticsearch": {"results": {}, "search_time_ms": 0, "total_matches": 0, "max_score": 0},
                "performance_comparison": {},
                "feature_comparison": {}
            }

        # Get advanced features demo (with fallback to empty results)
        try:
            advanced_features = await comparison_service.get_advanced_features_demo(
                user_id=str(user.id),
                required_skills=analysis_request.required_skills or []
            )
        except Exception as e:
            logger.warning(f"Advanced features demo failed, using empty results: {e}")
            advanced_features = {
                "fuzzy_matching": {"results": []},
                "synonym_matching": {"results": []},
                "aggregations": {"results": {}}
            }

        # Create analysis record
        analysis = ElasticJobAnalysis(
            user_id=str(user.id),
            job_description=job_description,
            job_url=analysis_request.job_url,
            chromadb_results=comparison_results["chromadb"].get("results", {}),
            chromadb_search_time_ms=comparison_results["chromadb"].get("search_time_ms"),
            chromadb_matches_count=comparison_results["chromadb"].get("total_matches"),
            chromadb_relevance_score=comparison_results["chromadb"].get("relevance_scores", [0])[0] if comparison_results["chromadb"].get("relevance_scores") else None,
            elasticsearch_results=comparison_results["elasticsearch"].get("results", {}),
            elasticsearch_search_time_ms=comparison_results["elasticsearch"].get("search_time_ms"),
            elasticsearch_matches_count=comparison_results["elasticsearch"].get("total_matches"),
            elasticsearch_relevance_score=comparison_results["elasticsearch"].get("max_score"),
            fuzzy_matches=advanced_features.get("fuzzy_matching", {}).get("results", []),
            synonym_matches=advanced_features.get("synonym_matching", {}).get("results", []),
            skill_clusters=advanced_features.get("aggregations", {}).get("results", {}),
            performance_comparison=comparison_results.get("performance_comparison", {}),
            feature_comparison=comparison_results.get("feature_comparison", {}),
            # LLM Analysis Results
            job_analysis=llm_analysis.get("job_analysis", {}),
            fit_score=llm_analysis.get("fit_score", {}),
            success_probability=llm_analysis.get("interview_success", llm_analysis.get("success_probability", {})),  # Use interview_success or fallback to success_probability
            interview_success=llm_analysis.get("interview_success"),  # New field
            provider=analysis_request.provider
        )

        db.add(analysis)
        await db.commit()
        await db.refresh(analysis)

        logger.info(f"Job analysis completed for user {user.id}, analysis ID {analysis.id}")
        return analysis

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing job: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/comparison/{analysis_id}", response_model=JobAnalysisResponse)
async def get_comparison(
    analysis_id: int,
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """Get specific analysis/comparison results by ID."""
    statement = select(ElasticJobAnalysis).where(
        ElasticJobAnalysis.id == analysis_id,
        ElasticJobAnalysis.user_id == str(user.id)
    )
    result = await db.execute(statement)
    analysis = result.scalars().first()

    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")

    return analysis


@router.get("/comparisons", response_model=List[JobAnalysisResponse])
async def list_comparisons(
    limit: int = Query(10, ge=1, le=50),
    offset: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """List all comparison analyses for current user."""
    statement = select(ElasticJobAnalysis).where(
        ElasticJobAnalysis.user_id == str(user.id)
    ).order_by(ElasticJobAnalysis.created_at.desc()).limit(limit).offset(offset)

    result = await db.execute(statement)
    analyses = result.scalars().all()

    return analyses


@router.get("/advanced-features")
async def get_advanced_features(
    user: User = Depends(current_active_user),
):
    """
    Get demonstration of Elasticsearch advanced features.

    Shows:
    - Fuzzy matching (typo tolerance)
    - Synonym support
    - Multi-field weighted search
    - Aggregations and analytics
    - Highlighting
    """
    # Check if profile exists
    # For demo purposes, we'll use common tech skills
    demo_skills = ["Python", "JavaScript", "Docker", "Kubernetes", "Machine Learning"]

    try:
        advanced_features = await comparison_service.get_advanced_features_demo(
            user_id=str(user.id),
            required_skills=demo_skills
        )
        return advanced_features
    except Exception as e:
        logger.error(f"Error getting advanced features: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/debug/cleanup-old-analyses")
async def cleanup_old_analyses(
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """Delete old analyses with wrong format (IDs 23, 24)."""
    try:
        from sqlalchemy import delete

        # Delete specific old analyses
        stmt = delete(ElasticJobAnalysis).where(
            ElasticJobAnalysis.user_id == str(user.id),
            ElasticJobAnalysis.id.in_([23, 24])
        )
        result = await db.execute(stmt)
        await db.commit()

        # Count remaining
        count_stmt = select(ElasticJobAnalysis).where(
            ElasticJobAnalysis.user_id == str(user.id)
        )
        count_result = await db.execute(count_stmt)
        remaining = len(count_result.scalars().all())

        return {
            "status": "success",
            "deleted": result.rowcount,
            "remaining_analyses": remaining
        }
    except Exception as e:
        import traceback
        return {
            "status": "error",
            "error": str(e),
            "traceback": traceback.format_exc()
        }


@router.post("/debug/simple-save")
async def simple_save_test(
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """Test endpoint - just save minimal data without LLM."""
    try:
        analysis = ElasticJobAnalysis(
            user_id=str(user.id),
            job_description="Test job",
            job_url=None,
            chromadb_results={},
            chromadb_search_time_ms=0,
            chromadb_matches_count=0,
            chromadb_relevance_score=None,
            elasticsearch_results={},
            elasticsearch_search_time_ms=0,
            elasticsearch_matches_count=0,
            elasticsearch_relevance_score=None,
            fuzzy_matches=[],
            synonym_matches=[],
            skill_clusters={},
            performance_comparison={},
            feature_comparison={},
            job_analysis={"test": "data"},
            fit_score={"total": 0},
            success_probability={"probability": 0},
            provider="test"
        )

        db.add(analysis)
        await db.commit()
        await db.refresh(analysis)

        return {
            "status": "success",
            "id": analysis.id,
            "message": "Data saved successfully"
        }
    except Exception as e:
        import traceback
        return {
            "status": "error",
            "error": str(e),
            "traceback": traceback.format_exc()
        }


@router.get("/debug/columns")
async def check_columns(db: AsyncSession = Depends(get_async_session)):
    """Debug endpoint to check if LLM analysis columns exist."""
    try:
        result = await db.execute(text("""
            SELECT column_name
            FROM information_schema.columns
            WHERE table_name = 'elastic_job_analyses'
            AND column_name IN ('job_analysis', 'fit_score', 'success_probability')
            ORDER BY column_name;
        """))
        columns = [row[0] for row in result]

        return {
            "table": "elastic_job_analyses",
            "columns_found": columns,
            "all_present": len(columns) == 3,
            "expected": ["fit_score", "job_analysis", "success_probability"]
        }
    except Exception as e:
        return {"error": str(e)}


@router.get("/health")
async def health_check():
    """Check Elasticsearch connection health."""
    try:
        # Try to ping Elasticsearch
        health = es_service.client.cluster.health()
        return {
            "status": "healthy",
            "elasticsearch": {
                "status": health["status"],
                "number_of_nodes": health["number_of_nodes"],
                "active_shards": health["active_shards"]
            }
        }
    except Exception as e:
        logger.error(f"Elasticsearch health check failed: {e}")
        return {
            "status": "unhealthy",
            "error": str(e)
        }


@router.post("/parse-doc")
async def parse_doc_file(file: UploadFile = File(...)):
    """
    Parse .doc or .docx file and extract text.

    This endpoint accepts both old .doc format and modern .docx format files.
    """
    try:
        import io

        # Read file contents
        file_content = await file.read()
        filename = file.filename.lower() if file.filename else ""

        # Handle .docx files
        if filename.endswith('.docx'):
            try:
                import docx
                from io import BytesIO

                # Parse DOCX file
                doc = docx.Document(BytesIO(file_content))

                # Extract all text from paragraphs
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)

                text = '\n'.join(text_parts)

                if text and len(text) > 10:
                    logger.info(f"Successfully parsed .docx file: {len(text)} characters")
                    return {"text": text, "success": True}
                else:
                    raise ValueError("Could not extract meaningful text from .docx file")

            except Exception as docx_err:
                logger.error(f"Error parsing .docx file: {docx_err}")
                raise HTTPException(
                    status_code=400,
                    detail=f".docx file parsing failed: {str(docx_err)}. The file may be corrupted. Please try copy/paste instead."
                )

        # Handle .doc files (legacy format)
        elif filename.endswith('.doc'):
            import olefile

            # Try to parse .doc file using olefile
            ole = olefile.OleFileIO(io.BytesIO(file_content))

            # Try to find WordDocument stream
            if ole.exists('WordDocument'):
                word_stream = ole.openstream('WordDocument')
                data = word_stream.read()

                # Extract text (very basic extraction - .doc format is complex)
                # This will get some text but may not be perfect
                text = data.decode('latin-1', errors='ignore')

                # Clean up non-printable characters
                import string
                printable = set(string.printable)
                text = ''.join(filter(lambda x: x in printable, text))

                ole.close()

                if text and len(text) > 10:
                    return {"text": text, "success": True}
                else:
                    raise ValueError("Could not extract meaningful text from .doc file")
            else:
                raise ValueError(".doc file structure not recognized")
        else:
            raise ValueError(f"Unsupported file type. Expected .doc or .docx, got: {filename}")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error parsing Word document: {e}")
        raise HTTPException(
            status_code=400,
            detail=f"Word document parsing failed: {str(e)}. Please save as .txt or copy/paste the text instead."
        )


@router.get("/fetch-url")
async def fetch_url_proxy(url: str = Query(..., description="URL to fetch")):
    """
    Proxy endpoint to fetch content from a URL.

    This bypasses CORS restrictions by fetching the URL from the backend.
    Used for loading job descriptions and CVs from external URLs.
    """
    try:
        import httpx

        # Validate URL
        if not url or not url.startswith(('http://', 'https://')):
            raise HTTPException(
                status_code=400,
                detail="Invalid URL. Must start with http:// or https://"
            )

        # Fetch the URL with timeout
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            response = await client.get(url)
            response.raise_for_status()

            # Return the text content
            return {"text": response.text, "success": True}

    except httpx.HTTPError as e:
        logger.error(f"Error fetching URL {url}: {e}")
        raise HTTPException(
            status_code=400,
            detail=f"Failed to fetch URL: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Unexpected error fetching URL {url}: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to fetch URL: {str(e)}"
        )


# ============================================================================
# Logstash Integration Endpoints
# ============================================================================

@router.post("/logstash/parse-cv")
async def logstash_parse_cv(
    cv_text: str,
    user: User = Depends(current_active_user),
):
    """
    Parse CV using Logstash pipeline (or simulated).

    Extracts:
    - Skills (programming languages, frameworks, databases, etc.)
    - Years of experience
    - Education level
    - Job titles

    Returns:
        Parsed CV data
    """
    try:
        result = logstash_service.parse_cv(cv_text, str(user.id))
        return {
            "status": "success",
            "data": result,
            "logstash_deployed": logstash_service.is_logstash_available
        }
    except Exception as e:
        logger.error(f"Error parsing CV: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/logstash/parse-job")
async def logstash_parse_job(
    job_description: str,
    job_id: str = "generated",
):
    """
    Parse job description using Logstash pipeline (or simulated).

    Extracts:
    - Company name
    - Location
    - Remote policy
    - Required skills
    - Years of experience required
    - Salary range

    Returns:
        Parsed job data
    """
    try:
        import uuid
        if job_id == "generated":
            job_id = str(uuid.uuid4())[:8]

        result = logstash_service.parse_job(job_description, job_id)
        return {
            "status": "success",
            "data": result,
            "logstash_deployed": logstash_service.is_logstash_available
        }
    except Exception as e:
        logger.error(f"Error parsing job: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/logstash/pipeline-status")
async def get_logstash_pipeline_status():
    """
    Get status of all Logstash pipelines.

    Returns:
        Pipeline status for cv-parsing, job-parsing, enrichment
    """
    try:
        status = logstash_service.get_pipeline_status()
        return status
    except Exception as e:
        logger.error(f"Error getting pipeline status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/elastic-stack/status")
async def get_elastic_stack_status():
    """
    Get status of entire Elastic Stack (Elasticsearch, Logstash, Kibana).

    Returns:
        Status of all Elastic Stack components
    """
    try:
        status = {
            "elasticsearch": {
                "available": es_service.is_available(),
                "url": os.getenv("ELASTICSEARCH_URL", "elasticsearch.railway.internal:9200"),
                "status": "active" if es_service.is_available() else "unavailable"
            },
            "logstash": {
                "available": logstash_service.is_logstash_available,
                "url": os.getenv("LOGSTASH_URL", "not_deployed"),
                "status": "deployed" if logstash_service.is_logstash_available else "simulated",
                "pipelines": logstash_service.get_pipeline_status()["pipelines"]
            },
            "kibana": {
                "available": os.getenv("KIBANA_URL") is not None,
                "url": os.getenv("KIBANA_URL", "not_deployed"),
                "status": "deployed" if os.getenv("KIBANA_URL") else "pending"
            },
            "stack_ready": es_service.is_available()  # At minimum, Elasticsearch must be available
        }

        return status
    except Exception as e:
        logger.error(f"Error getting Elastic Stack status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== Advanced Elasticsearch Features Endpoints =====

@router.post("/advanced/phrase-search")
async def elasticsearch_phrase_search(
    phrase: str,
    slop: int = 2,
    user: User = Depends(current_active_user),
):
    """
    Demonstrate phrase matching with configurable slop.

    Example: "Senior Python Developer" with slop=2 matches "Senior Software Python Developer"
    """
    if not es_service.is_available():
        raise HTTPException(status_code=503, detail="Elasticsearch not available")

    try:
        results = await es_service.phrase_search(
            phrase=phrase,
            user_id=str(user.id),
            slop=slop
        )

        return {
            "status": "success",
            "feature": "phrase_matching",
            "description": f"Exact phrase search with slop={slop} (allows {slop} words in between)",
            "example": "Finds 'Senior Python Developer' even if other words appear between terms",
            "benefit": "More flexible than exact match, stricter than full-text search",
            "data": results
        }
    except Exception as e:
        logger.error(f"Phrase search error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/advanced/wildcard-search")
async def elasticsearch_wildcard_search(
    pattern: str,
    field: str = "skills",
    user: User = Depends(current_active_user),
):
    """
    Demonstrate wildcard pattern matching.

    Examples:
    - "Java*" matches Java, JavaScript, JavaFX
    - "Pyth?n" matches Python, Pythyn
    - "*Script" matches JavaScript, TypeScript
    """
    if not es_service.is_available():
        raise HTTPException(status_code=503, detail="Elasticsearch not available")

    try:
        results = await es_service.wildcard_search(
            pattern=pattern,
            field=field,
            user_id=str(user.id)
        )

        return {
            "status": "success",
            "feature": "wildcard_search",
            "description": "Pattern matching with * (any characters) and ? (single character)",
            "examples": [
                "Java* → Java, JavaScript, JavaFX",
                "Pyth?n → Python, Pythyn",
                "*Script → JavaScript, TypeScript"
            ],
            "benefit": "Find variations of technologies without knowing exact names",
            "data": results
        }
    except Exception as e:
        logger.error(f"Wildcard search error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/advanced/aggregations")
async def elasticsearch_advanced_aggregations(
    user: User = Depends(current_active_user),
):
    """
    Demonstrate advanced aggregations: experience buckets, skill distribution, statistics.
    """
    if not es_service.is_available():
        raise HTTPException(status_code=503, detail="Elasticsearch not available")

    try:
        results = await es_service.get_advanced_aggregations(user_id=str(user.id))

        return {
            "status": "success",
            "feature": "advanced_aggregations",
            "description": "Real-time analytics on CV data",
            "capabilities": [
                "Experience level bucketing (Junior, Mid, Senior, Expert)",
                "Education distribution analysis",
                "Top skills ranking",
                "Experience statistics (min, max, avg, median)",
                "Skill co-occurrence by experience level"
            ],
            "benefit": "Instant insights without separate analytics database",
            "data": results
        }
    except Exception as e:
        logger.error(f"Advanced aggregations error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/advanced/explain-match")
async def elasticsearch_explain_match(
    job_description: str,
    required_skills: Optional[List[str]] = None,
    user: User = Depends(current_active_user),
):
    """
    Explain why a CV matched a job (or didn't) with score breakdown.

    Uses Elasticsearch Explain API to show scoring details.
    """
    if not es_service.is_available():
        raise HTTPException(status_code=503, detail="Elasticsearch not available")

    try:
        explanation = await es_service.explain_match(
            user_id=str(user.id),
            job_description=job_description,
            required_skills=required_skills or []
        )

        return {
            "status": "success",
            "feature": "explain_api",
            "description": "Detailed scoring explanation showing why CV matched job",
            "use_cases": [
                "Debug why matches appear in certain order",
                "Understand which skills contributed most to score",
                "Optimize CV based on scoring factors",
                "Transparency for candidates"
            ],
            "benefit": "Complete transparency in search relevance",
            "data": explanation
        }
    except Exception as e:
        logger.error(f"Explain match error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/advanced/search-suggestions")
async def elasticsearch_search_suggestions(
    prefix: str,
    field: str = "skills",
    size: int = 10
):
    """
    Get autocomplete suggestions for search terms.

    Example: prefix="Py" → ["Python", "PyTorch", "Pydantic"]
    """
    if not es_service.is_available():
        raise HTTPException(status_code=503, detail="Elasticsearch not available")

    try:
        suggestions = await es_service.get_search_suggestions(
            prefix=prefix,
            field=field,
            size=size
        )

        return {
            "status": "success",
            "feature": "autocomplete",
            "description": "Real-time search suggestions as user types",
            "example": f"'{prefix}' → {suggestions[:3]}",
            "benefit": "Better UX, helps users discover available skills",
            "data": suggestions
        }
    except Exception as e:
        logger.error(f"Search suggestions error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/advanced/multi-index-search")
async def elasticsearch_multi_index_search(
    query_text: str,
):
    """
    Search across CV and Job indices simultaneously.

    Demonstrates Elasticsearch ability to search multiple indices in single request.
    """
    if not es_service.is_available():
        raise HTTPException(status_code=503, detail="Elasticsearch not available")

    try:
        results = await es_service.multi_index_search(query_text=query_text)

        return {
            "status": "success",
            "feature": "multi_index_search",
            "description": "Search multiple data types simultaneously",
            "example": "Search 'Python Django' across CVs and Job postings in one query",
            "benefit": "Faster than separate queries, unified relevance scoring",
            "data": results
        }
    except Exception as e:
        logger.error(f"Multi-index search error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/advanced/features-overview")
async def get_elasticsearch_features_overview():
    """
    Get comprehensive overview of all Elasticsearch advanced features.

    Use this to showcase Elasticsearch capabilities vs ChromaDB.
    """
    return {
        "status": "success",
        "elastic_stack_advantages": {
            "elasticsearch": {
                "features": [
                    {
                        "name": "Fuzzy Matching",
                        "description": "Handles typos and misspellings (e.g., 'Pythn' → 'Python')",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/advanced/fuzzy-matches"
                    },
                    {
                        "name": "Synonym Search",
                        "description": "Matches related terms (e.g., 'ML' → 'Machine Learning')",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/advanced/synonym-matches"
                    },
                    {
                        "name": "Phrase Matching",
                        "description": "Find exact phrases with flexible word order",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/advanced/phrase-search"
                    },
                    {
                        "name": "Wildcard Search",
                        "description": "Pattern matching (Java* → Java, JavaScript)",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/advanced/wildcard-search"
                    },
                    {
                        "name": "Real-time Aggregations",
                        "description": "Analytics on skills, experience, education",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/advanced/aggregations"
                    },
                    {
                        "name": "Explain API",
                        "description": "Shows why a match scored the way it did",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/advanced/explain-match"
                    },
                    {
                        "name": "Autocomplete",
                        "description": "Search suggestions as user types",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/advanced/search-suggestions"
                    },
                    {
                        "name": "Multi-Index Search",
                        "description": "Search CVs and Jobs in single request",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/advanced/multi-index-search"
                    },
                    {
                        "name": "Weighted Multi-Field Search",
                        "description": "Skills weighted 3x, CV text 2x, job titles 1.5x",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/analyze"
                    },
                    {
                        "name": "Highlighting",
                        "description": "Shows which parts of CV matched the query",
                        "chromadb_has": False,
                        "endpoint": "/elasticsearch/analyze"
                    }
                ],
                "performance": {
                    "scalability": "Horizontal scaling to petabytes",
                    "speed": "Sub-millisecond search on millions of documents",
                    "concurrency": "Handles thousands of concurrent queries"
                }
            },
            "logstash": {
                "features": [
                    {
                        "name": "CV Parsing Pipeline",
                        "description": "Extract skills, experience, education from CV",
                        "endpoint": "/elasticsearch/logstash/parse-cv"
                    },
                    {
                        "name": "Job Parsing Pipeline",
                        "description": "Extract requirements, salary, location from job",
                        "endpoint": "/elasticsearch/logstash/parse-job"
                    },
                    {
                        "name": "Data Enrichment",
                        "description": "Add synonyms, categorization, metadata",
                        "endpoint": "/elasticsearch/logstash/parse-cv"
                    },
                    {
                        "name": "Pipeline Monitoring",
                        "description": "Health checks and throughput metrics",
                        "endpoint": "/elasticsearch/logstash/pipeline-status"
                    }
                ],
                "benefits": [
                    "Process 1000s of CVs per minute",
                    "Standardize unstructured data",
                    "Real-time data transformation",
                    "100+ built-in filters and processors"
                ]
            },
            "kibana": {
                "features": [
                    "Interactive dashboards",
                    "Real-time analytics",
                    "Custom visualizations",
                    "Index management",
                    "Query builder UI"
                ],
                "note": "Custom visualizations implemented with Chart.js/D3.js"
            },
            "chromadb_limitations": [
                "No fuzzy matching - exact or semantic only",
                "No synonym support",
                "Limited query types (embedding similarity)",
                "No aggregations or analytics",
                "No explain/debugging capabilities",
                "No autocomplete support",
                "Single collection per query",
                "No phrase or wildcard search"
            ]
        },
        "integration_status": {
            "elasticsearch": es_service.is_available(),
            "logstash": logstash_service.is_logstash_available,
            "kibana": "Not deployed yet (see ELASTIC_STACK_DEPLOYMENT.md)"
        }
    }


# ===== Demo Data Generator Endpoints =====

@router.post("/demo/generate")
async def generate_demo_data(
    count: int = 50,
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """
    Generate demo CV profile data for Analytics Dashboard.

    Creates synthetic CV chunks with skills, databases, programming languages,
    companies, and certifications. Data is indexed in Elasticsearch to populate
    the Analytics Dashboard visualizations.
    """
    try:
        import random

        logger.info(f"Generating {count} demo CV profiles for user {user.id}")

        # Sample data pools for generating realistic profiles
        skills_pool = [
            "Python", "JavaScript", "TypeScript", "React", "Vue.js", "Angular",
            "Node.js", "Django", "FastAPI", "Flask", "Express.js", "Docker",
            "Kubernetes", "AWS", "Azure", "GCP", "PostgreSQL", "MongoDB",
            "Redis", "Elasticsearch", "GraphQL", "REST API", "Git", "CI/CD",
            "Machine Learning", "Data Analysis", "SQL", "NoSQL", "Microservices",
            "Agile", "Scrum", "TDD", "Linux", "Bash", "Java", "C++", "Go",
            "Rust", "Ruby", "PHP", "Swift", "Kotlin"
        ]

        databases_pool = [
            "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch",
            "Oracle", "SQL Server", "DynamoDB", "Cassandra", "Neo4j",
            "MariaDB", "SQLite", "CouchDB", "InfluxDB", "TimescaleDB"
        ]

        languages_pool = [
            "Python", "JavaScript", "TypeScript", "Java", "C++", "C#",
            "Go", "Rust", "Ruby", "PHP", "Swift", "Kotlin", "Scala",
            "R", "MATLAB", "Perl", "Shell", "PowerShell"
        ]

        companies_pool = [
            "Google", "Amazon", "Microsoft", "Meta", "Apple", "Netflix",
            "Tesla", "SpaceX", "IBM", "Oracle", "SAP", "Salesforce",
            "Adobe", "Intel", "NVIDIA", "Uber", "Airbnb", "Stripe",
            "Shopify", "Spotify", "Twitter", "LinkedIn", "GitHub",
            "Atlassian", "Dropbox", "Zoom", "Slack", "DocuSign"
        ]

        certifications_pool = [
            "AWS Certified Solutions Architect", "Azure Administrator",
            "Google Cloud Professional", "Kubernetes Administrator (CKA)",
            "Docker Certified Associate", "Certified Scrum Master",
            "PMP", "CISSP", "CompTIA Security+", "Oracle Certified Professional"
        ]

        user_id = str(user.id)
        index_name = f"cv_showcase_{user_id}"

        # Ensure index exists
        if not es_service.client.indices.exists(index=index_name):
            logger.info(f"Creating index {index_name}")
            es_service.client.indices.create(
                index=index_name,
                body={
                    "mappings": {
                        "properties": {
                            "content": {"type": "text"},
                            "chunk_id": {"type": "integer"},
                            "embedding": {"type": "dense_vector", "dims": 384},
                            "skills": {"type": "keyword"},
                            "databases": {"type": "keyword"},
                            "programming_languages": {"type": "keyword"},
                            "companies": {"type": "keyword"},
                            "certifications": {"type": "keyword"}
                        }
                    }
                }
            )

        # Generate and index demo chunks
        profiles_created = 0
        for i in range(count):
            # Random selection of skills, databases, etc.
            num_skills = random.randint(5, 15)
            num_dbs = random.randint(2, 6)
            num_langs = random.randint(2, 5)
            num_companies = random.randint(1, 4)
            num_certs = random.randint(0, 3)

            selected_skills = random.sample(skills_pool, num_skills)
            selected_dbs = random.sample(databases_pool, num_dbs)
            selected_langs = random.sample(languages_pool, num_langs)
            selected_companies = random.sample(companies_pool, num_companies)
            selected_certs = random.sample(certifications_pool, num_certs) if num_certs > 0 else []

            # Generate realistic CV text chunk
            experience_years = random.randint(1, 15)
            content = f"Senior developer with {experience_years} years of experience. "
            content += f"Worked at {', '.join(selected_companies[:2])}. "
            content += f"Expert in {', '.join(selected_skills[:5])}. "
            content += f"Proficient with {', '.join(selected_dbs[:3])} databases. "
            content += f"Programming languages: {', '.join(selected_langs)}. "
            if selected_certs:
                content += f"Certifications: {', '.join(selected_certs)}."

            # Create document for Elasticsearch
            doc = {
                "content": content,
                "chunk_id": i,
                "embedding": [random.random() for _ in range(384)],  # Dummy embedding
                "skills": selected_skills,
                "databases": selected_dbs,
                "programming_languages": selected_langs,
                "companies": selected_companies,
                "certifications": selected_certs
            }

            # Index document
            es_service.client.index(
                index=index_name,
                id=f"demo_{i}",
                body=doc
            )
            profiles_created += 1

            if profiles_created % 10 == 0:
                logger.info(f"Indexed {profiles_created}/{count} demo profiles")

        # Refresh index to make data searchable
        es_service.client.indices.refresh(index=index_name)

        logger.info(f"Successfully generated {profiles_created} demo profiles")

        return {
            "status": "success",
            "message": f"Generated {profiles_created} demo CV profiles",
            "profiles_created": profiles_created,
            "index_name": index_name,
            "next_steps": [
                "View Analytics Dashboard to see visualizations",
                "Data includes skills, databases, programming languages, companies, and certifications"
            ]
        }

    except Exception as e:
        logger.error(f"Error generating demo data: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/demo/stats")
async def get_demo_statistics(
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """
    Get statistics on demo data for visualizations.

    Returns aggregated data perfect for charts:
    - Fit score distribution
    - Success probability ranges
    - Performance comparison stats
    - Skill frequency
    - Company distribution
    """
    try:
        # Get all analyses for user
        statement = select(ElasticJobAnalysis).where(
            ElasticJobAnalysis.user_id == str(user.id)
        )
        result = await db.execute(statement)
        analyses = result.scalars().all()

        if not analyses:
            return {
                "status": "success",
                "message": "No data available. Generate demo data first.",
                "total_analyses": 0
            }

        # Calculate statistics
        fit_scores = [a.fit_score.get("total", 0) for a in analyses if a.fit_score]
        success_probs = [a.success_probability.get("probability", 0) for a in analyses if a.success_probability]

        # Performance stats
        chromadb_times = [a.chromadb_search_time_ms for a in analyses if a.chromadb_search_time_ms]
        es_times = [a.elasticsearch_search_time_ms for a in analyses if a.elasticsearch_search_time_ms]

        # Skill extraction
        all_skills = {}
        for a in analyses:
            if a.job_analysis and "requirements" in a.job_analysis:
                must_have = a.job_analysis["requirements"].get("must_have", [])
                for skill in must_have:
                    all_skills[skill] = all_skills.get(skill, 0) + 1

        top_skills = sorted(all_skills.items(), key=lambda x: x[1], reverse=True)[:20]

        # Company distribution
        companies = {}
        for a in analyses:
            if a.job_analysis:
                company = a.job_analysis.get("company")
                if company:
                    companies[company] = companies.get(company, 0) + 1

        return {
            "status": "success",
            "total_analyses": len(analyses),
            "fit_score_distribution": {
                "min": min(fit_scores) if fit_scores else 0,
                "max": max(fit_scores) if fit_scores else 0,
                "avg": sum(fit_scores) / len(fit_scores) if fit_scores else 0,
                "buckets": {
                    "0-20": len([s for s in fit_scores if 0 <= s < 20]),
                    "20-40": len([s for s in fit_scores if 20 <= s < 40]),
                    "40-60": len([s for s in fit_scores if 40 <= s < 60]),
                    "60-80": len([s for s in fit_scores if 60 <= s < 80]),
                    "80-100": len([s for s in fit_scores if 80 <= s <= 100])
                }
            },
            "success_probability_distribution": {
                "min": min(success_probs) if success_probs else 0,
                "max": max(success_probs) if success_probs else 0,
                "avg": sum(success_probs) / len(success_probs) if success_probs else 0,
                "buckets": {
                    "Low (0-40)": len([p for p in success_probs if 0 <= p < 40]),
                    "Medium (40-70)": len([p for p in success_probs if 40 <= p < 70]),
                    "High (70-100)": len([p for p in success_probs if 70 <= p <= 100])
                }
            },
            "performance_comparison": {
                "chromadb": {
                    "avg_time_ms": sum(chromadb_times) / len(chromadb_times) if chromadb_times else 0,
                    "min_time_ms": min(chromadb_times) if chromadb_times else 0,
                    "max_time_ms": max(chromadb_times) if chromadb_times else 0
                },
                "elasticsearch": {
                    "avg_time_ms": sum(es_times) / len(es_times) if es_times else 0,
                    "min_time_ms": min(es_times) if es_times else 0,
                    "max_time_ms": max(es_times) if es_times else 0
                },
                "speedup_factor": (sum(chromadb_times) / len(chromadb_times)) / (sum(es_times) / len(es_times)) if chromadb_times and es_times else 0
            },
            "top_skills": [{"skill": skill, "count": count} for skill, count in top_skills],
            "company_distribution": [{"company": comp, "count": count} for comp, count in sorted(companies.items(), key=lambda x: x[1], reverse=True)[:10]],
            "visualization_ready": True,
            "chart_recommendations": [
                "Bar chart: Fit score distribution",
                "Pie chart: Success probability ranges",
                "Line chart: Performance comparison (ChromaDB vs Elasticsearch)",
                "Word cloud: Top skills",
                "Horizontal bar: Company distribution"
            ]
        }

    except Exception as e:
        logger.error(f"Error getting demo statistics: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/demo/clear")
async def clear_demo_data(
    keep_latest: int = 1,
    db: AsyncSession = Depends(get_async_session),
    user: User = Depends(current_active_user),
):
    """
    Clear demo data, optionally keeping the N most recent analyses.

    Args:
        keep_latest: Number of recent analyses to keep (default: 1)
    """
    try:
        from sqlalchemy import delete

        # Get IDs to keep
        if keep_latest > 0:
            statement = select(ElasticJobAnalysis.id).where(
                ElasticJobAnalysis.user_id == str(user.id)
            ).order_by(ElasticJobAnalysis.created_at.desc()).limit(keep_latest)

            result = await db.execute(statement)
            keep_ids = [row[0] for row in result]

            # Delete all except those to keep
            stmt = delete(ElasticJobAnalysis).where(
                ElasticJobAnalysis.user_id == str(user.id),
                ElasticJobAnalysis.id.notin_(keep_ids)
            )
        else:
            # Delete all
            stmt = delete(ElasticJobAnalysis).where(
                ElasticJobAnalysis.user_id == str(user.id)
            )

        result = await db.execute(stmt)
        await db.commit()

        return {
            "status": "success",
            "deleted_count": result.rowcount,
            "kept_latest": keep_latest
        }

    except Exception as e:
        logger.error(f"Error clearing demo data: {e}")
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Database Visualization Endpoints
# ============================================================================

@router.get("/database/stats")
async def get_database_stats(
    user: User = Depends(current_active_user),
    db: AsyncSession = Depends(get_async_session)
):
    """Get database statistics for ChromaDB and Elasticsearch."""
    try:
        # Get all profiles for this user
        profile_stmt = select(UserElasticProfile).where(
            UserElasticProfile.user_id == str(user.id)
        )
        profile_result = await db.execute(profile_stmt)
        profiles = profile_result.scalars().all()

        # Get all job analyses for this user
        analysis_stmt = select(ElasticJobAnalysis).where(
            ElasticJobAnalysis.user_id == str(user.id)
        )
        analysis_result = await db.execute(analysis_stmt)
        analyses = analysis_result.scalars().all()

        # Calculate statistics
        total_profiles = len(profiles)
        total_jobs = len(analyses)

        # Calculate average skills per profile
        total_skills_chromadb = 0
        total_skills_elastic = 0

        for profile in profiles:
            if profile.skills_extracted:
                skills_list = profile.skills_extracted if isinstance(profile.skills_extracted, list) else json.loads(profile.skills_extracted)
                total_skills_chromadb += len(skills_list)
                total_skills_elastic += len(skills_list)

        avg_skills_chromadb = total_skills_chromadb / total_profiles if total_profiles > 0 else 0
        avg_skills_elastic = total_skills_elastic / total_profiles if total_profiles > 0 else 0

        # Total documents = profiles + job analyses
        total_docs_chromadb = total_profiles + total_jobs
        total_docs_elastic = total_profiles + total_jobs

        return {
            "chromadb": {
                "total_documents": total_docs_chromadb,
                "total_profiles": total_profiles,
                "total_jobs": total_jobs,
                "avg_skills_per_profile": round(avg_skills_chromadb, 1)
            },
            "elasticsearch": {
                "total_documents": total_docs_elastic,
                "total_profiles": total_profiles,
                "total_jobs": total_jobs,
                "avg_skills_per_profile": round(avg_skills_elastic, 1)
            }
        }

    except Exception as e:
        logger.error(f"Error getting database stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/chromadb/documents")
async def get_chromadb_documents(
    user: User = Depends(current_active_user),
    db: AsyncSession = Depends(get_async_session),
    limit: int = Query(100, ge=1, le=1000)
):
    """Get ChromaDB documents (profiles and job analyses)."""
    try:
        # Get profiles
        profile_stmt = select(UserElasticProfile).where(
            UserElasticProfile.user_id == str(user.id)
        ).limit(limit)
        profile_result = await db.execute(profile_stmt)
        profiles = profile_result.scalars().all()

        # Get job analyses
        analysis_stmt = select(ElasticJobAnalysis).where(
            ElasticJobAnalysis.user_id == str(user.id)
        ).limit(limit)
        analysis_result = await db.execute(analysis_stmt)
        analyses = analysis_result.scalars().all()

        documents = []

        # Add profiles
        for profile in profiles:
            documents.append({
                "type": "profile",
                "id": profile.id,
                "created_at": profile.created_at.isoformat() if profile.created_at else None,
                "skills": profile.skills_extracted if isinstance(profile.skills_extracted, list) else json.loads(profile.skills_extracted) if profile.skills_extracted else [],
                "experience_years": profile.experience_years,
                "education_level": profile.education_level,
            })

        # Add job analyses
        for analysis in analyses:
            documents.append({
                "type": "job_analysis",
                "id": analysis.id,
                "created_at": analysis.created_at.isoformat() if analysis.created_at else None,
                "job_title": getattr(analysis, 'job_title', None),
                "chromadb_search_time_ms": analysis.chromadb_search_time_ms,
                "elasticsearch_search_time_ms": analysis.elasticsearch_search_time_ms,
            })

        return {
            "total": len(documents),
            "documents": documents
        }

    except Exception as e:
        logger.error(f"Error getting ChromaDB documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/elasticsearch/documents")
async def get_elasticsearch_documents(
    user: User = Depends(current_active_user),
    db: AsyncSession = Depends(get_async_session),
    limit: int = Query(100, ge=1, le=1000)
):
    """Get Elasticsearch documents (profiles and job analyses)."""
    try:
        # Get profiles
        profile_stmt = select(UserElasticProfile).where(
            UserElasticProfile.user_id == str(user.id)
        ).limit(limit)
        profile_result = await db.execute(profile_stmt)
        profiles = profile_result.scalars().all()

        # Get job analyses
        analysis_stmt = select(ElasticJobAnalysis).where(
            ElasticJobAnalysis.user_id == str(user.id)
        ).limit(limit)
        analysis_result = await db.execute(analysis_stmt)
        analyses = analysis_result.scalars().all()

        documents = []

        # Add profiles
        for profile in profiles:
            documents.append({
                "type": "profile",
                "id": profile.id,
                "created_at": profile.created_at.isoformat() if profile.created_at else None,
                "skills": profile.skills_extracted if isinstance(profile.skills_extracted, list) else json.loads(profile.skills_extracted) if profile.skills_extracted else [],
                "experience_years": profile.experience_years,
                "education_level": profile.education_level,
            })

        # Add job analyses
        for analysis in analyses:
            documents.append({
                "type": "job_analysis",
                "id": analysis.id,
                "created_at": analysis.created_at.isoformat() if analysis.created_at else None,
                "job_title": getattr(analysis, 'job_title', None),
                "chromadb_search_time_ms": analysis.chromadb_search_time_ms,
                "elasticsearch_search_time_ms": analysis.elasticsearch_search_time_ms,
            })

        return {
            "total": len(documents),
            "documents": documents
        }

    except Exception as e:
        logger.error(f"Error getting Elasticsearch documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# RAG COMPARISON SHOWCASE ENDPOINT
# ============================================================================

# 5 Example queries for RAG comparison
EXAMPLE_QUERIES = [
    "What experience does Michael Dabrock have at Google?",
    "Which vector database does he currently use in his projects?",
    "Explain the GDPR advantages of local LLMs in his applications.",
    "What role did he play at Cognizant and IBM?",
    "How does he integrate RAG in tools like CV Matcher?"
]

@router.get("/rag-comparison")
async def rag_comparison_demo(
    llm: str = Query("local", description="LLM provider: 'grok' or 'local'"),
    user: User = Depends(current_active_user),
    db: AsyncSession = Depends(get_async_session),
):
    """
    RAG Comparison Showcase: Compare ChromaDB vs Elasticsearch with 5 example queries.

    This endpoint demonstrates the quality difference between:
    - ChromaDB: Pure vector similarity search
    - Elasticsearch: Hybrid search (BM25 + kNN) with fuzzy matching, synonyms, and field boosting

    For each query, we:
    1. Retrieve top-3 chunks from both databases
    2. Generate an answer using the selected LLM (Grok or Local Ollama)
    3. Compare retrieval scores and answer quality

    Returns side-by-side comparison results with summary metrics.
    """
    try:
        logger.info(f"RAG Comparison requested with LLM: {llm}, User: {user.id}")

        # Validate LLM parameter
        if llm not in ["grok", "local"]:
            raise HTTPException(status_code=400, detail="Invalid LLM. Must be 'grok' or 'local'")

        # Get user's CV data for RAG
        profile = await db.execute(
            select(UserElasticProfile).where(UserElasticProfile.user_id == str(user.id))
        )
        user_profile = profile.scalar_one_or_none()

        if not user_profile:
            raise HTTPException(
                status_code=404,
                detail="No profile found. Please upload your CV first in the 'Analyze Job' tab."
            )

        # Run all 5 queries in parallel
        import asyncio
        import time

        results = []
        for query in EXAMPLE_QUERIES:
            logger.info(f"Processing query: {query}")

            # Run pgvector and Elasticsearch RAG in parallel
            chromadb_task = asyncio.create_task(
                run_chromadb_rag(query, user.id, llm, db)
            )
            elasticsearch_task = asyncio.create_task(
                run_elasticsearch_rag(query, user.id, llm, db)
            )

            chromadb_result, elasticsearch_result = await asyncio.gather(
                chromadb_task, elasticsearch_task
            )

            # Compare results
            score_delta = elasticsearch_result["avg_score"] - chromadb_result["avg_score"]
            if score_delta > 0.05:
                winner = "elasticsearch"
            elif score_delta < -0.05:
                winner = "chromadb"
            else:
                winner = "tie"

            results.append({
                "query": query,
                "chromadb": chromadb_result,
                "elasticsearch": elasticsearch_result,
                "winner": winner,
                "score_delta": round(score_delta, 3)
            })

        # Compute summary metrics
        avg_chromadb_score = sum(r["chromadb"]["avg_score"] for r in results) / len(results)
        avg_elasticsearch_score = sum(r["elasticsearch"]["avg_score"] for r in results) / len(results)

        elasticsearch_wins = sum(1 for r in results if r["winner"] == "elasticsearch")
        chromadb_wins = sum(1 for r in results if r["winner"] == "chromadb")
        ties = sum(1 for r in results if r["winner"] == "tie")

        summary = {
            "avg_chromadb_score": round(avg_chromadb_score, 3),
            "avg_elasticsearch_score": round(avg_elasticsearch_score, 3),
            "elasticsearch_win_rate": round(elasticsearch_wins / len(results), 2),
            "total_queries": len(results),
            "elasticsearch_wins": elasticsearch_wins,
            "chromadb_wins": chromadb_wins,
            "ties": ties
        }

        logger.info(f"RAG Comparison completed. Summary: {summary}")

        return {
            "llm_used": llm,
            "queries": results,
            "summary": summary
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in RAG comparison: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


async def run_chromadb_rag(query: str, user_id, llm: str, db: AsyncSession):
    """Run RAG pipeline with pgvector (pure vector similarity) - formerly ChromaDB"""
    import time
    start_time = time.time()

    try:
        # Retrieve from pgvector (vector similarity only)
        if not vector_service.is_available():
            # Return mock data if pgvector not available
            return {
                "answer": "pgvector not available. Embedding model not loaded.",
                "chunks": [],
                "retrieval_time_ms": 0,
                "avg_score": 0.0
            }

        # Search pgvector
        from uuid import UUID
        search_results = await vector_service.query(
            session=db,
            user_id=UUID(str(user_id)),
            query_text=query,
            project_id=None,
            n_results=3
        )

        retrieval_time = (time.time() - start_time) * 1000

        if not search_results:
            return {
                "answer": "No relevant information found in pgvector.",
                "chunks": [],
                "retrieval_time_ms": round(retrieval_time, 2),
                "avg_score": 0.0
            }

        # Format chunks for LLM
        chunks_text = "\n\n".join([
            f"[Source: {r['metadata'].get('type', 'unknown')}]\n{r['content']}"
            for r in search_results
        ])

        # Generate answer with LLM
        prompt = f"""Based on the following context, answer the question concisely and accurately.

Context:
{chunks_text}

Question: {query}

Answer:"""

        answer = await generate_llm_answer(prompt, llm)

        # Calculate average score
        scores = [r.get('score', 0.5) for r in search_results]
        avg_score = sum(scores) / len(scores) if scores else 0.0

        # Format chunks for response
        chunks = [
            {
                "text": r['content'][:500],  # Truncate long chunks
                "source": r['metadata'].get('type', 'cv.pdf'),
                "score": round(r.get('score', 0.5), 3)
            }
            for r in search_results
        ]

        return {
            "answer": answer,
            "chunks": chunks,
            "retrieval_time_ms": round(retrieval_time, 2),
            "avg_score": round(avg_score, 3)
        }

    except Exception as e:
        logger.error(f"ChromaDB RAG error: {e}")
        return {
            "answer": f"Error: {str(e)}",
            "chunks": [],
            "retrieval_time_ms": 0,
            "avg_score": 0.0
        }


async def run_elasticsearch_rag(query: str, user_id, llm: str, db: AsyncSession):
    """Run RAG pipeline with Elasticsearch (hybrid BM25 + kNN)"""
    import time
    start_time = time.time()

    try:
        # Hybrid search in Elasticsearch
        if not es_service.is_available():
            return {
                "answer": "Elasticsearch not available. Please check configuration.",
                "chunks": [],
                "retrieval_time_ms": 0,
                "avg_score": 0.0
            }

        # Search Elasticsearch with hybrid approach
        es_results = await es_service.hybrid_search(
            query=query,
            user_id=str(user_id),
            top_k=3
        )

        retrieval_time = (time.time() - start_time) * 1000

        if not es_results:
            return {
                "answer": "No relevant information found in Elasticsearch.",
                "chunks": [],
                "retrieval_time_ms": round(retrieval_time, 2),
                "avg_score": 0.0
            }

        # Format chunks for LLM
        chunks_text = "\n\n".join([
            f"[Source: {r.get('source', 'cv.pdf')}]\n{r.get('text', '')}"
            for r in es_results
        ])

        # Generate answer with LLM
        prompt = f"""Based on the following context, answer the question concisely and accurately.

Context:
{chunks_text}

Question: {query}

Answer:"""

        answer = await generate_llm_answer(prompt, llm)

        # Calculate average score
        scores = [r.get('score', 0.7) for r in es_results]
        avg_score = sum(scores) / len(scores) if scores else 0.0

        # Format chunks for response
        chunks = [
            {
                "text": r.get('text', '')[:500],  # Truncate long chunks
                "source": r.get('source', 'cv.pdf'),
                "score": round(r.get('score', 0.7), 3)
            }
            for r in es_results
        ]

        return {
            "answer": answer,
            "chunks": chunks,
            "retrieval_time_ms": round(retrieval_time, 2),
            "avg_score": round(avg_score, 3)
        }

    except Exception as e:
        logger.error(f"Elasticsearch RAG error: {e}")
        return {
            "answer": f"Error: {str(e)}",
            "chunks": [],
            "retrieval_time_ms": 0,
            "avg_score": 0.0
        }


async def generate_llm_answer(prompt: str, llm_provider: str) -> str:
    """Generate answer using selected LLM (Grok or Local Ollama)"""
    try:
        if llm_provider == "grok":
            # Use Grok API via LLMGateway
            response = await llm_gateway.generate(
                prompt=prompt,
                provider="grok",
                temperature=0.3,
                max_tokens=200
            )
            return response.get("content", "No answer generated.")

        elif llm_provider == "local":
            # Use Ollama on Railway (private network)
            import httpx
            from backend.config import settings

            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.post(
                    f"{settings.OLLAMA_BASE_URL}/api/generate",
                    json={
                        "model": settings.OLLAMA_MODEL,
                        "prompt": prompt,
                        "stream": False,
                        "options": {
                            "temperature": 0.3,
                            "num_predict": 200
                        }
                    }
                )

                if response.status_code == 200:
                    result = response.json()
                    return result.get("response", "No answer generated.")
                else:
                    logger.error(f"Ollama error: {response.status_code}")
                    return "Ollama service unavailable. Please ensure Ollama is running."

        else:
            return "Invalid LLM provider."

    except Exception as e:
        logger.error(f"LLM generation error: {e}")
        return f"Error generating answer: {str(e)}"


@router.get("/database-stats")
async def get_database_stats(
    db: AsyncSession = Depends(get_async_session)
):
    """Get statistics about vector database entries."""
    try:
        stats = {
            "pgvector": {
                "available": vector_service.is_available(),
                "count": 0
            },
            "elasticsearch": {
                "available": es_service.is_available(),
                "count": 0
            },
            "ollama_model": os.getenv("OLLAMA_MODEL", "llama3.2:3b")
        }

        # Get pgvector count from Document table
        if vector_service.is_available():
            try:
                from backend.models.document import Document, DocumentType
                # Count all CV_SHOWCASE documents
                result = await db.execute(
                    select(Document).where(Document.type == DocumentType.CV_SHOWCASE)
                )
                documents = result.scalars().all()
                stats["pgvector"]["count"] = len(documents)
            except Exception as e:
                # Silently skip if CV_SHOWCASE enum not in database yet
                if "CV_SHOWCASE" in str(e) and "enum" in str(e).lower():
                    logger.warning("⚠️  pgvector CV_SHOWCASE enum not available - count set to 0")
                else:
                    logger.error(f"Failed to get pgvector count: {e}")
                stats["pgvector"]["count"] = 0

        # Get Elasticsearch count
        if es_service.is_available():
            try:
                # Count documents in CV index (synchronous call)
                result = es_service.client.count(index=es_service.cv_index)
                stats["elasticsearch"]["count"] = result.get("count", 0)
            except Exception as e:
                logger.error(f"Failed to get Elasticsearch count: {e}")

        return stats
    except Exception as e:
        logger.error(f"Error getting database stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/enum-diagnostic")
async def check_enum_status(
    db: AsyncSession = Depends(get_async_session),
    current_user: User = Depends(current_active_user)
):
    """Diagnostic endpoint to check if CV_SHOWCASE enum exists"""
    try:
        from backend.models.document import Document, DocumentType

        # Try to query with CV_SHOWCASE enum
        try:
            result = await db.execute(
                select(Document).where(Document.type == DocumentType.CV_SHOWCASE).limit(1)
            )
            docs = result.scalars().all()

            return {
                "enum_exists": True,
                "enum_value": DocumentType.CV_SHOWCASE.value,
                "document_count": len(docs),
                "message": "✅ CV_SHOWCASE enum exists in database"
            }
        except Exception as enum_error:
            return {
                "enum_exists": False,
                "error": str(enum_error),
                "error_type": type(enum_error).__name__,
                "message": "❌ CV_SHOWCASE enum does NOT exist in database"
            }
    except Exception as e:
        logger.error(f"Enum diagnostic failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/test-pgvector")
async def test_pgvector_indexing(
    db: AsyncSession = Depends(get_async_session),
    current_user: User = Depends(current_active_user)
):
    """Diagnostic endpoint to test pgvector indexing and return actual error"""
    try:
        from backend.services.elasticsearch_vector_service import ElasticsearchVectorService
        from uuid import UUID

        test_vector_service = ElasticsearchVectorService()

        # Try to add a test document (skip is_available check to avoid caching issues)
        test_content = "This is a test CV content for pgvector diagnostic testing."
        test_metadata = {"type": "cv", "skills": "Python,FastAPI", "job_titles": "Developer"}

        logger.info("🔍 Testing pgvector add_documents directly...")
        try:
            chunks = await test_vector_service.add_documents(
                session=db,
                user_id=UUID(str(current_user.id)),
                documents=[{
                    "id": f"test_cv_{current_user.id}",
                    "content": test_content,
                    "metadata": test_metadata
                }],
                project_id=None,
                chunk_size=500
            )

            return {
                "success": True,
                "chunks_added": chunks,
                "message": f"✅ pgvector test succeeded - {chunks} chunks added"
            }
        except Exception as add_error:
            import traceback
            return {
                "success": False,
                "error": str(add_error),
                "error_type": type(add_error).__name__,
                "traceback": traceback.format_exc(),
                "message": "❌ pgvector add_documents failed"
            }

    except Exception as e:
        import traceback
        logger.error(f"pgvector test failed: {e}")
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__,
            "traceback": traceback.format_exc(),
            "message": "❌ pgvector test endpoint failed"
        }


@router.get("/check-enum-values")
async def check_enum_values(
    current_user: User = Depends(current_active_user),
    db: AsyncSession = Depends(get_async_session)
):
    """Query PostgreSQL to see what enum values actually exist"""
    try:
        from sqlalchemy import text

        # Use AsyncSession - clean async approach without raw_connection()
        result = await db.execute(text(
            """
            SELECT enumlabel
            FROM pg_enum
            WHERE enumtypid = 'documenttype'::regtype
            ORDER BY enumsortorder
            """
        ))
        enum_values = [row[0] for row in result.fetchall()]

        return {
            "enum_name": "documenttype",
            "values": enum_values,
            "count": len(enum_values),
            "has_cv_showcase": "cv_showcase" in enum_values,
            "has_CV_SHOWCASE": "CV_SHOWCASE" in enum_values,
            "message": f"✅ Found {len(enum_values)} enum values: {enum_values}"
        }

    except Exception as e:
        logger.error(f"Failed to query enum values: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to query enum: {str(e)}")


@router.post("/fix-enum")
async def fix_enum_value(
    db: AsyncSession = Depends(get_async_session),
    current_user: User = Depends(current_active_user)
):
    """Emergency fix: Manually add CV_SHOWCASE enum value to database

    CRITICAL: ALTER TYPE ADD VALUE cannot run in a transaction block!
    Runs in background thread with autocommit to bypass transaction.
    """
    try:
        import asyncio
        from concurrent.futures import ThreadPoolExecutor
        from sqlalchemy import create_engine, text
        from backend.config import settings

        logger.info("🔧 Manually adding CV_SHOWCASE enum value to database...")
        logger.info("⚠️  Using sync engine in thread pool (ADD VALUE cannot run in transactions)")

        def add_enum_sync():
            """Run ALTER TYPE in sync context with autocommit using raw psycopg2"""
            import psycopg2
            from urllib.parse import urlparse

            # Parse DATABASE_URL
            db_url = settings.DATABASE_URL.replace('postgresql+asyncpg://', 'postgresql://')
            parsed = urlparse(db_url)

            # Connect with psycopg2 (autocommit mode)
            conn = psycopg2.connect(
                host=parsed.hostname,
                port=parsed.port or 5432,
                user=parsed.username,
                password=parsed.password,
                database=parsed.path.lstrip('/')
            )
            conn.set_session(autocommit=True)

            try:
                cur = conn.cursor()
                cur.execute("ALTER TYPE documenttype ADD VALUE IF NOT EXISTS 'CV_SHOWCASE'")
                cur.close()
            finally:
                conn.close()

        # Run in thread pool to avoid blocking async loop
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as executor:
            await loop.run_in_executor(executor, add_enum_sync)

        logger.info("✅ CV_SHOWCASE enum value added successfully (uppercase to match SQL Enum names)")

        # Enum successfully added - connection pool refresh needed
        return {
            "success": True,
            "message": "✅ CV_SHOWCASE enum value added successfully (uppercase)",
            "enum_value": "CV_SHOWCASE",
            "method": "Sync engine with AUTOCOMMIT in thread pool",
            "note": "Added 'CV_SHOWCASE' (uppercase) to match SQLAlchemy's enum name convention. Restart the application for all connections to see it."
        }

    except Exception as e:
        logger.error(f"Failed to add CV_SHOWCASE enum: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to add enum: {str(e)}")


@router.get("/current-model")
async def get_current_model(
    current_user: User = Depends(current_active_user),
):
    """Get the currently configured LLM model."""
    try:
        ollama_model = os.getenv("OLLAMA_MODEL", "llama3.2:3b")

        return {
            "model": ollama_model,
            "provider": "local",
            "type": "ollama",
            "description": f"Local Ollama model: {ollama_model}"
        }
    except Exception as e:
        logger.error(f"Failed to get current model: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get model: {str(e)}")


@router.get("/recreate-indices-public")
async def recreate_elasticsearch_indices_public():
    """Public endpoint to delete and recreate Elasticsearch indices with correct analyzer configuration."""
    try:
        if not es_service.is_available():
            return {"error": "Elasticsearch not available"}

        # Delete existing indices
        deleted = []
        if es_service.client.indices.exists(index=es_service.cv_index):
            es_service.client.indices.delete(index=es_service.cv_index)
            deleted.append(es_service.cv_index)
            logger.info(f"Deleted index: {es_service.cv_index}")

        if es_service.client.indices.exists(index=es_service.job_index):
            es_service.client.indices.delete(index=es_service.job_index)
            deleted.append(es_service.job_index)
            logger.info(f"Deleted index: {es_service.job_index}")

        # Recreate indices with correct configuration
        es_service._ensure_indices()
        logger.info("Recreated indices with skill_analyzer")

        return {
            "status": "success",
            "deleted_indices": deleted,
            "message": "Indices recreated successfully with skill_analyzer. You need to re-import your CV data now."
        }

    except Exception as e:
        logger.error(f"Failed to recreate indices: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Recreate failed: {str(e)}")


@router.get("/recreate-indices")
async def recreate_elasticsearch_indices(
    current_user: User = Depends(current_active_user),
):
    """Delete and recreate Elasticsearch indices with correct analyzer configuration."""
    try:
        if not es_service.is_available():
            return {"error": "Elasticsearch not available"}

        # Delete existing indices
        deleted = []
        if es_service.client.indices.exists(index=es_service.cv_index):
            es_service.client.indices.delete(index=es_service.cv_index)
            deleted.append(es_service.cv_index)
            logger.info(f"Deleted index: {es_service.cv_index}")

        if es_service.client.indices.exists(index=es_service.job_index):
            es_service.client.indices.delete(index=es_service.job_index)
            deleted.append(es_service.job_index)
            logger.info(f"Deleted index: {es_service.job_index}")

        # Recreate indices with correct configuration
        es_service._ensure_indices()
        logger.info("Recreated indices with skill_analyzer")

        return {
            "status": "success",
            "deleted_indices": deleted,
            "message": "Indices recreated successfully with skill_analyzer. You need to re-import your CV data now."
        }

    except Exception as e:
        logger.error(f"Failed to recreate indices: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Recreate failed: {str(e)}")


@router.get("/debug-index")
async def debug_elasticsearch_index(
    current_user: User = Depends(current_active_user),
):
    """Debug endpoint to see what's actually indexed in Elasticsearch."""
    try:
        if not es_service.is_available():
            return {"error": "Elasticsearch not available"}

        # Get index mapping
        mapping = es_service.client.indices.get_mapping(index=es_service.cv_index)

        # Get a sample document for this user
        search_result = es_service.client.search(
            index=es_service.cv_index,
            body={
                "query": {
                    "term": {"user_id": str(current_user.id)}
                },
                "size": 1
            }
        )

        sample_doc = None
        if search_result["hits"]["total"]["value"] > 0:
            sample_doc = search_result["hits"]["hits"][0]["_source"]
            # Truncate long fields for readability
            if "cv_text" in sample_doc and len(sample_doc["cv_text"]) > 500:
                sample_doc["cv_text"] = sample_doc["cv_text"][:500] + "... (truncated)"

        # Get count
        count_result = es_service.client.count(
            index=es_service.cv_index,
            body={"query": {"term": {"user_id": str(current_user.id)}}}
        )

        return {
            "index_name": es_service.cv_index,
            "total_documents": count_result["count"],
            "mapping": mapping,
            "sample_document": sample_doc,
            "field_names": list(sample_doc.keys()) if sample_doc else []
        }

    except Exception as e:
        logger.error(f"Debug index failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Debug failed: {str(e)}")


@router.post("/compare-query")
async def compare_query(
    question: str = Query(..., description="Question to ask both vector databases"),
    provider: str = Query("local", description="LLM provider: 'local' or 'grok'"),
    db: AsyncSession = Depends(get_async_session),
    current_user: User = Depends(current_active_user),
):
    """Compare a single query across pgvector and Elasticsearch with LLM evaluation.

    This endpoint:
    1. Searches both pgvector and Elasticsearch for relevant chunks
    2. Generates answers using the selected LLM
    3. Uses LLM to evaluate which answer is better
    4. Returns structured comparison with winner highlighted
    """
    try:
        import time

        # Map frontend provider names to backend provider names
        provider_mapping = {
            "local": "ollama",
            "grok": "grok",
            "anthropic": "anthropic"
        }
        llm_provider = provider_mapping.get(provider.lower(), "grok")

        logger.info(f"🔍 Comparing query across vector DBs: '{question}' (provider={provider} -> {llm_provider})")

        # Step 1: Search pgvector
        logger.info("📊 Searching pgvector...")
        pgvector_start = time.time()
        pgvector_chunks = await vector_service.query(
            session=db,
            user_id=current_user.id,
            query_text=question,
            n_results=3
        )
        pgvector_time = (time.time() - pgvector_start) * 1000

        # Step 2: Search Elasticsearch
        logger.info("🔍 Searching Elasticsearch...")
        es_start = time.time()
        es_chunks = await es_service.hybrid_search(
            user_id=str(current_user.id),
            query=question,
            top_k=3
        )
        es_time = (time.time() - es_start) * 1000

        # Step 3: Generate answer from pgvector chunks
        logger.info("🤖 Generating pgvector answer...")
        pgvector_context = "\n\n".join([
            f"[Chunk {i+1}] {chunk.get('content', chunk.get('text', ''))}"
            for i, chunk in enumerate(pgvector_chunks)
        ])

        pgvector_prompt = f"""Answer the following question based ONLY on the context provided. Be CONCISE and PRECISE.

RULES:
- Maximum 2-3 sentences
- Use concrete facts and numbers from the context
- No filler words or unnecessary explanations
- If the answer is a list, use bullet points

Question: {question}

Context:
{pgvector_context}

Answer:"""

        pgvector_response = llm_gateway.generate(
            prompt=pgvector_prompt,
            provider=llm_provider,
            temperature=0.3,
            max_tokens=150  # Reduced from 500 for concise answers
        )
        pgvector_answer = pgvector_response.get('response', '')

        # Step 4: Generate answer from Elasticsearch chunks
        logger.info("🤖 Generating Elasticsearch answer...")
        es_context = "\n\n".join([
            f"[Chunk {i+1}] {chunk.get('content', chunk.get('_source', {}).get('content', ''))}"
            for i, chunk in enumerate(es_chunks)
        ])

        es_prompt = f"""Answer the following question based ONLY on the context provided. Be CONCISE and PRECISE.

RULES:
- Maximum 2-3 sentences
- Use concrete facts and numbers from the context
- No filler words or unnecessary explanations
- If the answer is a list, use bullet points

Question: {question}

Context:
{es_context}

Answer:"""

        es_response = llm_gateway.generate(
            prompt=es_prompt,
            provider=llm_provider,
            temperature=0.3,
            max_tokens=150  # Reduced from 500 for concise answers
        )
        es_answer = es_response.get('response', '')

        # Step 5: RAGAS-based evaluation of answers
        logger.info(f"⚖️  Evaluating answers with RAGAS (provider={llm_provider})...")

        # Import RAGAS evaluator (uses same provider as answer generation for DSGVO compliance)
        from backend.services.ragas_evaluator import get_ragas_evaluator
        ragas = get_ragas_evaluator(provider=llm_provider)

        # Extract chunk texts for RAGAS
        pgvector_chunk_texts = [
            chunk.get('content', chunk.get('text', ''))
            for chunk in pgvector_chunks
        ]
        es_chunk_texts = [
            chunk.get('content', chunk.get('text', ''))
            for chunk in es_chunks
        ]

        # Run RAGAS comparison
        ragas_comparison = ragas.compare_systems(
            question=question,
            answer_a=pgvector_answer,
            answer_b=es_answer,
            contexts_a=pgvector_chunk_texts,
            contexts_b=es_chunk_texts,
            ground_truth=None,  # Can be added later for specific benchmark questions
            system_a_name="pgvector",
            system_b_name="elasticsearch"
        )

        # Convert RAGAS scores (0-1) to percentage (0-100)
        pgvector_score = int(ragas_comparison['pgvector']['overall_score'] * 100)
        elasticsearch_score = int(ragas_comparison['elasticsearch']['overall_score'] * 100)

        # Build evaluation result
        evaluation = {
            "winner": ragas_comparison['winner'],
            "reasoning": ragas_comparison.get('reasoning', ''),
            "pgvector_score": pgvector_score,
            "elasticsearch_score": elasticsearch_score,
            "ragas_metrics": {
                "pgvector": ragas_comparison['pgvector']['scores'],
                "elasticsearch": ragas_comparison['elasticsearch']['scores']
            },
            "evaluation_method": "ragas_grok",
            "confidence": ragas_comparison.get('confidence', 'medium')
        }

        # Format chunks for response
        pgvector_chunks_formatted = [
            {
                "text": chunk.get('content', chunk.get('text', '')),
                "source": chunk.get('metadata', {}).get('source', 'pgvector'),
                "score": chunk.get('score', chunk.get('distance', 0.0))
            }
            for chunk in pgvector_chunks[:3]
        ]

        es_chunks_formatted = [
            {
                "text": chunk.get('content', chunk.get('_source', {}).get('content', '')),
                "source": chunk.get('_source', {}).get('metadata', {}).get('source', 'elasticsearch'),
                "score": chunk.get('_score', 0.0) / 10.0  # Normalize ES score to 0-1 range
            }
            for chunk in es_chunks[:3]
        ]

        logger.info(f"✅ Comparison complete! Winner: {evaluation['winner']}")

        return {
            "question": question,
            "pgvector": {
                "answer": pgvector_answer,
                "chunks": pgvector_chunks_formatted,
                "retrieval_time_ms": pgvector_time,
                "score": evaluation["pgvector_score"]
            },
            "elasticsearch": {
                "answer": es_answer,
                "chunks": es_chunks_formatted,
                "retrieval_time_ms": es_time,
                "score": evaluation["elasticsearch_score"]
            },
            "evaluation": {
                "winner": evaluation["winner"],
                "reasoning": evaluation["reasoning"],
                "pgvector_score": evaluation["pgvector_score"],
                "elasticsearch_score": evaluation["elasticsearch_score"]
            },
            "llm_used": provider,
            "timestamp": datetime.utcnow().isoformat()
        }

    except Exception as e:
        logger.error(f"Failed to compare query: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Comparison failed: {str(e)}")


@router.get("/aggregations")
async def get_elasticsearch_aggregations(current_user: User = Depends(current_active_user)):
    """
    Get aggregated statistics from Elasticsearch for analytics dashboard.
    Returns top databases, programming languages, companies, certifications, and skills.
    """
    try:
        user_id = str(current_user.id)
        index_name = f"cv_showcase_{user_id}"

        # Check if index exists
        if not es_service.client.indices.exists(index=index_name):
            return {
                "databases": [],
                "programming_languages": [],
                "companies": [],
                "certifications": [],
                "skills": [],
                "total_chunks": 0
            }

        # Build aggregation query
        agg_query = {
            "size": 0,  # We only want aggregations, not documents
            "aggs": {
                "databases": {
                    "terms": {
                        "field": "databases.keyword",
                        "size": 20
                    }
                },
                "programming_languages": {
                    "terms": {
                        "field": "programming_languages.keyword",
                        "size": 20
                    }
                },
                "companies": {
                    "terms": {
                        "field": "companies.keyword",
                        "size": 20
                    }
                },
                "certifications": {
                    "terms": {
                        "field": "certifications.keyword",
                        "size": 20
                    }
                },
                "skills": {
                    "terms": {
                        "field": "skills.keyword",
                        "size": 30
                    }
                }
            }
        }

        # Execute aggregation query
        response = es_service.client.search(index=index_name, **agg_query)

        # Format results
        return {
            "databases": [
                {"name": bucket["key"], "count": bucket["doc_count"]}
                for bucket in response["aggregations"]["databases"]["buckets"]
            ],
            "programming_languages": [
                {"name": bucket["key"], "count": bucket["doc_count"]}
                for bucket in response["aggregations"]["programming_languages"]["buckets"]
            ],
            "companies": [
                {"name": bucket["key"], "count": bucket["doc_count"]}
                for bucket in response["aggregations"]["companies"]["buckets"]
            ],
            "certifications": [
                {"name": bucket["key"], "count": bucket["doc_count"]}
                for bucket in response["aggregations"]["certifications"]["buckets"]
            ],
            "skills": [
                {"name": bucket["key"], "count": bucket["doc_count"]}
                for bucket in response["aggregations"]["skills"]["buckets"]
            ],
            "total_chunks": response["hits"]["total"]["value"]
        }

    except Exception as e:
        logger.error(f"Failed to get aggregations: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Aggregation failed: {str(e)}")


@router.post("/faceted-search")
async def faceted_search(
    query: str = "",
    databases: List[str] = [],
    programming_languages: List[str] = [],
    companies: List[str] = [],
    certifications: List[str] = [],
    current_user: User = Depends(current_active_user)
):
    """
    Perform faceted search with filters.
    Combines Elasticsearch hybrid search with filtering by databases, programming languages, companies, and certifications.
    """
    try:
        user_id = str(current_user.id)
        index_name = f"cv_showcase_{user_id}"

        # Check if index exists
        if not es_service.client.indices.exists(index=index_name):
            return {
                "results": [],
                "total": 0,
                "query": query,
                "filters": {
                    "databases": databases,
                    "programming_languages": programming_languages,
                    "companies": companies,
                    "certifications": certifications
                }
            }

        # Build filter clauses
        filter_clauses = []

        if databases:
            filter_clauses.append({"terms": {"databases.keyword": databases}})

        if programming_languages:
            filter_clauses.append({"terms": {"programming_languages.keyword": programming_languages}})

        if companies:
            filter_clauses.append({"terms": {"companies.keyword": companies}})

        if certifications:
            filter_clauses.append({"terms": {"certifications.keyword": certifications}})

        # Build search query
        if query:
            # If there's a search query, use hybrid search
            # Generate embedding for the query
            embedding_response = requests.post(
                "http://localhost:11434/api/embeddings",
                json={"model": "nomic-embed-text", "prompt": query}
            )
            query_embedding = embedding_response.json()["embedding"]

            # Build hybrid search with filters
            search_query = {
                "query": {
                    "bool": {
                        "must": [
                            {
                                "multi_match": {
                                    "query": query,
                                    "fields": [
                                        "content^1",
                                        "databases^5",
                                        "programming_languages^4",
                                        "companies^3",
                                        "certifications^2",
                                        "skills^2"
                                    ],
                                    "fuzziness": "AUTO"
                                }
                            }
                        ],
                        "filter": filter_clauses
                    }
                },
                "knn": {
                    "field": "embedding",
                    "query_vector": query_embedding,
                    "k": 10,
                    "num_candidates": 100,
                    "filter": {
                        "bool": {
                            "filter": filter_clauses
                        }
                    } if filter_clauses else None
                },
                "size": 20,
                "rank": {
                    "rrf": {}
                }
            }
        else:
            # If no search query, just filter
            search_query = {
                "query": {
                    "bool": {
                        "filter": filter_clauses if filter_clauses else [{"match_all": {}}]
                    }
                },
                "size": 20
            }

        # Execute search
        response = es_service.client.search(index=index_name, **search_query)

        # Format results
        results = []
        for hit in response["hits"]["hits"]:
            source = hit["_source"]
            results.append({
                "content": source.get("content", ""),
                "databases": source.get("databases", []),
                "programming_languages": source.get("programming_languages", []),
                "companies": source.get("companies", []),
                "certifications": source.get("certifications", []),
                "skills": source.get("skills", []),
                "score": hit["_score"]
            })

        return {
            "results": results,
            "total": response["hits"]["total"]["value"],
            "query": query,
            "filters": {
                "databases": databases,
                "programming_languages": programming_languages,
                "companies": companies,
                "certifications": certifications
            }
        }

    except Exception as e:
        logger.error(f"Failed faceted search: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Faceted search failed: {str(e)}")


@router.get("/analytics")
async def get_analytics_data(current_user: User = Depends(current_active_user)):
    """
    Get analytics data for the dashboard.
    Returns performance metrics, query distribution, and data insights.
    """
    try:
        user_id = str(current_user.id)
        index_name = f"cv_showcase_{user_id}"

        # Check if index exists
        if not es_service.client.indices.exists(index=index_name):
            return {
                "total_documents": 0,
                "index_size_bytes": 0,
                "avg_chunk_size": 0,
                "field_coverage": {},
                "top_skills": [],
                "timeline": [],
                "database_distribution": [],
                "language_distribution": []
            }

        # Get index stats
        stats = es_service.client.indices.stats(index=index_name)
        total_docs = stats["indices"][index_name]["total"]["docs"]["count"]
        index_size = stats["indices"][index_name]["total"]["store"]["size_in_bytes"]

        # Get all documents to calculate metrics
        all_docs_query = {
            "query": {"match_all": {}},
            "size": 1000,  # Adjust if needed
            "_source": ["content", "databases", "programming_languages", "companies", "skills", "certifications"]
        }
        all_docs = es_service.client.search(index=index_name, **all_docs_query)

        # Calculate average chunk size
        total_chars = sum(len(doc["_source"].get("content", "")) for doc in all_docs["hits"]["hits"])
        avg_chunk_size = total_chars / total_docs if total_docs > 0 else 0

        # Calculate field coverage
        field_coverage = {
            "databases": sum(1 for doc in all_docs["hits"]["hits"] if doc["_source"].get("databases")) / total_docs * 100 if total_docs > 0 else 0,
            "programming_languages": sum(1 for doc in all_docs["hits"]["hits"] if doc["_source"].get("programming_languages")) / total_docs * 100 if total_docs > 0 else 0,
            "companies": sum(1 for doc in all_docs["hits"]["hits"] if doc["_source"].get("companies")) / total_docs * 100 if total_docs > 0 else 0,
            "skills": sum(1 for doc in all_docs["hits"]["hits"] if doc["_source"].get("skills")) / total_docs * 100 if total_docs > 0 else 0,
            "certifications": sum(1 for doc in all_docs["hits"]["hits"] if doc["_source"].get("certifications")) / total_docs * 100 if total_docs > 0 else 0
        }

        # Get top skills aggregation
        skills_agg = es_service.client.search(
            index=index_name,
            size=0,
            aggs={
                "top_skills": {
                    "terms": {
                        "field": "skills.keyword",
                        "size": 10
                    }
                }
            }
        )
        top_skills = [
            {"skill": bucket["key"], "count": bucket["doc_count"]}
            for bucket in skills_agg["aggregations"]["top_skills"]["buckets"]
        ]

        # Get database distribution
        db_agg = es_service.client.search(
            index=index_name,
            size=0,
            aggs={
                "databases": {
                    "terms": {
                        "field": "databases.keyword",
                        "size": 10
                    }
                }
            }
        )
        database_distribution = [
            {"name": bucket["key"], "value": bucket["doc_count"]}
            for bucket in db_agg["aggregations"]["databases"]["buckets"]
        ]

        # Get programming language distribution
        lang_agg = es_service.client.search(
            index=index_name,
            size=0,
            aggs={
                "languages": {
                    "terms": {
                        "field": "programming_languages.keyword",
                        "size": 10
                    }
                }
            }
        )
        language_distribution = [
            {"name": bucket["key"], "value": bucket["doc_count"]}
            for bucket in lang_agg["aggregations"]["languages"]["buckets"]
        ]

        # Create timeline (company work periods) - simplified for now
        company_agg = es_service.client.search(
            index=index_name,
            size=0,
            aggs={
                "companies": {
                    "terms": {
                        "field": "companies.keyword",
                        "size": 10
                    }
                }
            }
        )
        timeline = [
            {"company": bucket["key"], "mentions": bucket["doc_count"]}
            for bucket in company_agg["aggregations"]["companies"]["buckets"]
        ]

        return {
            "total_documents": total_docs,
            "index_size_bytes": index_size,
            "index_size_mb": round(index_size / (1024 * 1024), 2),
            "avg_chunk_size": round(avg_chunk_size, 0),
            "field_coverage": field_coverage,
            "top_skills": top_skills,
            "timeline": timeline,
            "database_distribution": database_distribution,
            "language_distribution": language_distribution
        }

    except Exception as e:
        logger.error(f"Failed to get analytics: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Analytics failed: {str(e)}")
