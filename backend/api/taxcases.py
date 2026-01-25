"""Tax Case Management API Endpoints"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import func, desc
from typing import List, Optional, Dict, Any
from datetime import datetime
from uuid import UUID
import os
import logging
import json
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

logger = logging.getLogger(__name__)

from backend.database import get_db
from backend.models.taxcase import TaxCase, TaxCaseDocument, TaxCaseFolder, TaxCaseExtractedData
from backend.models.user import User
from backend.schemas.taxcase import (
    TaxCaseCreate, TaxCaseResponse, TaxCaseDetailResponse,
    DocumentResponse, ExtractedDataResponse, ExtractedDataBulkUpdate,
    AuthRegisterRequest, AuthLoginRequest, AuthResponse
)
try:
    from backend.services.application_service import DocumentParser
except ImportError:
    DocumentParser = None

router = APIRouter()

# TEMPORARY: Fixed demo user UUID (same as applications)
DEMO_USER_ID = UUID("00000000-0000-0000-0000-000000000001")

def get_demo_user(db: Session = Depends(get_db)) -> User:
    """Return demo user from DB"""
    user = db.query(User).filter(User.id == DEMO_USER_ID).first()
    if not user:
        raise HTTPException(status_code=500, detail="Demo user not found - backend not initialized")
    return user


# Authentication Endpoints (Simplified for MVP - just return token)
@router.post("/auth/register", response_model=AuthResponse)
async def register_user(request: AuthRegisterRequest, db: Session = Depends(get_db)):
    """Register a new user with email (password sent via email)"""
    # For MVP: Just accept any email and return success
    token = f"tax_demo_token_{datetime.now().timestamp()}"

    logger.info(f"User registered (MVP mode): {request.email}")

    return AuthResponse(
        token=token,
        email=request.email,
        message=f"Registration successful. Password 'test123' sent to {request.email}"
    )


@router.post("/auth/login", response_model=AuthResponse)
async def login_user(request: AuthLoginRequest, db: Session = Depends(get_db)):
    """Login with email and password"""
    # For MVP: Accept any email with password "test123"
    if request.password != "test123":
        raise HTTPException(status_code=401, detail="Invalid email or password")

    token = f"tax_demo_token_{datetime.now().timestamp()}"

    logger.info(f"User logged in (MVP mode): {request.email}")

    return AuthResponse(
        token=token,
        email=request.email,
        message="Login successful"
    )


# Tax Case Endpoints
@router.get("/cases", response_model=List[TaxCaseResponse])
async def list_tax_cases(db: Session = Depends(get_db), user: User = Depends(get_demo_user)):
    """List all tax cases for the authenticated user"""
    cases = db.query(TaxCase).filter(TaxCase.user_id == user.id).order_by(desc(TaxCase.created_at)).all()

    response = []
    for case in cases:
        doc_count = db.query(func.count(TaxCaseDocument.id)).filter(
            TaxCaseDocument.tax_case_id == case.id
        ).scalar()

        response.append(TaxCaseResponse(
            id=case.id,
            name=case.name,
            status=case.status,
            validated=case.validated,
            document_count=doc_count or 0,
            created_at=case.created_at
        ))

    return response


@router.post("/cases", response_model=TaxCaseResponse)
async def create_tax_case(
    case_data: TaxCaseCreate,
    db: Session = Depends(get_db),
    user: User = Depends(get_demo_user)
):
    """Create a new tax case"""
    new_case = TaxCase(
        user_id=user.id,
        name=case_data.name,
        notes=case_data.notes,
        status="created"
    )

    db.add(new_case)
    db.commit()
    db.refresh(new_case)

    logger.info(f"Created tax case: {new_case.id} - {new_case.name}")

    return TaxCaseResponse(
        id=new_case.id,
        name=new_case.name,
        status=new_case.status,
        validated=new_case.validated,
        document_count=0,
        created_at=new_case.created_at
    )


@router.get("/cases/{case_id}", response_model=TaxCaseDetailResponse)
async def get_tax_case(case_id: int, db: Session = Depends(get_db), user: User = Depends(get_demo_user)):
    """Get detailed information about a tax case"""
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Tax case not found")

    # Get document count
    doc_count = db.query(func.count(TaxCaseDocument.id)).filter(
        TaxCaseDocument.tax_case_id == case.id
    ).scalar()

    # Get extracted data as dictionary
    extracted_items = db.query(TaxCaseExtractedData).filter(
        TaxCaseExtractedData.tax_case_id == case.id
    ).all()

    extracted_data = {}
    for item in extracted_items:
        extracted_data[item.field_name] = item.field_value

    return TaxCaseDetailResponse(
        id=case.id,
        name=case.name,
        status=case.status,
        validated=case.validated,
        notes=case.notes,
        document_count=doc_count or 0,
        extracted_data=extracted_data,
        created_at=case.created_at,
        updated_at=case.updated_at
    )


@router.post("/cases/{case_id}/upload")
async def upload_documents(
    case_id: int,
    files: List[UploadFile] = File(...),
    use_local_llm: str = Form("true"),  # DSGVO mode by default
    ocr_engine: str = Form("tesseract"),
    db: Session = Depends(get_db),
    user: User = Depends(get_demo_user)
):
    """Upload and process documents for a tax case"""
    prefer_local = use_local_llm.lower() == "true"
    ocr_mode = "PaddleOCR" if ocr_engine.lower() == "paddle" else "Tesseract OCR"
    logger.info(f"Upload mode: {'DSGVO (lokales LLM)' if prefer_local else 'Grok API'} mit {ocr_mode}")
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Tax case not found")

    # Create uploads directory if it doesn't exist
    upload_dir = f"/tmp/tax_uploads/{case_id}"
    os.makedirs(upload_dir, exist_ok=True)

    uploaded_docs = []

    try:
        for file in files:
            # Save file
            file_path = os.path.join(upload_dir, file.filename)
            with open(file_path, "wb") as f:
                content = await file.read()
                f.write(content)

            # For MVP: Just extract text from PDF or use placeholder
            doc_content = f"Document content from {file.filename}"

            # Try to parse with DocumentParser if available
            if DocumentParser:
                try:
                    parser = DocumentParser()
                    doc_content = parser.parse(file_path, ocr_engine=ocr_engine.lower())
                    logger.info(f"{ocr_mode} extracted text from {file.filename}")
                except Exception as parse_error:
                    logger.warning(f"{ocr_mode} could not parse document {file.filename}: {parse_error}")
                    # Use basic content
                    doc_content = f"Uploaded: {file.filename} ({ocr_mode} nicht verfügbar)"
            else:
                doc_content = f"Uploaded: {file.filename}"

            # Create document record
            doc = TaxCaseDocument(
                tax_case_id=case.id,
                filename=file.filename,
                file_path=file_path,
                doc_type="document",
                content=doc_content,
                validated=False
            )

            db.add(doc)
            db.flush()

            # Extract data using LLM (always creates data, even if LLM fails)
            await extract_data_from_document(case.id, doc.id, doc_content, db, prefer_local)

            uploaded_docs.append(doc.id)

        # Update case status
        case.status = "processing"
        db.commit()

        logger.info(f"Uploaded {len(files)} documents to case {case_id}")

        return {
            "success": True,
            "message": f"{len(files)} documents uploaded and processed",
            "document_ids": uploaded_docs
        }
    except Exception as e:
        db.rollback()
        logger.error(f"Error uploading documents: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


async def extract_data_from_document(case_id: int, doc_id: int, content: str, db: Session, prefer_local: bool = True):
    """Extract structured data from document using LLM"""

    # For MVP: Create sample extracted data
    llm_mode = "Lokales LLM (DSGVO)" if prefer_local else "Grok API"
    sample_data = {
        "document_content": content[:200] if content else "No content",
        "status": "Hochgeladen und bereit zur Verarbeitung",
        "verarbeitungsmodus": llm_mode,
        "hinweis": "Bitte bearbeiten Sie diese Felder und bestätigen Sie den Fall"
    }

    # Try to use LLM if available
    try:
        from backend.services.llm_gateway import llm_gateway

        # Prepare extraction prompt (H7-Formular für Kanarische Inseln)
        prompt = f"""Extrahiere folgende Informationen aus diesem Dokument für Import auf die Kanarischen Inseln:

1. SENDUNGSDATEN: art_der_sendung, warenwert_gesamt_eur, versandkosten, art_der_lieferung
2. ABSENDER: absender_name, absender_strasse, absender_plz, absender_ort, absender_land, absender_email, absender_telefon
3. EMPFÄNGER: empfaenger_name, empfaenger_strasse, empfaenger_plz, empfaenger_ort, empfaenger_insel, empfaenger_nif_nie_cif, empfaenger_email, empfaenger_telefon
4. WARENPOSITION: position_1_beschreibung, position_1_anzahl, position_1_stueckpreis, position_1_gesamtwert, position_1_ursprungsland
5. RECHNUNG: rechnungsnummer, rechnungsdatum, mehrwertsteuer_ausgewiesen

Dokumenteninhalt:
{content[:4000]}

Gib die Daten als JSON zurück. Verwende exakt die angegebenen Feldnamen.
Falls Felder nicht gefunden werden, lass sie weg.
"""

        # Use LLM to extract data with DSGVO preference
        provider = "ollama" if prefer_local else "grok"
        timeout = 240 if provider == "ollama" else 60  # 240 seconds for Ollama, 60 seconds for Grok
        result_dict = llm_gateway.generate(
            prompt=prompt,
            provider=provider,
            max_tokens=3000,
            timeout=timeout
        )
        result = result_dict.get("response", "")

        logger.info(f"LLM extraction mode: {llm_mode} using {provider}")

        # Parse JSON response
        try:
            # Use LLM gateway's robust JSON parser (handles markdown code blocks)
            extracted = llm_gateway.parse_json_response(result)
            # Flatten nested objects
            flattened = {}
            for key, value in extracted.items():
                if isinstance(value, dict):
                    for nested_key, nested_value in value.items():
                        flattened[nested_key] = nested_value
                else:
                    flattened[key] = value
            # Merge with sample data
            sample_data.update(flattened)
        except (json.JSONDecodeError, ValueError) as parse_error:
            # If parsing fails, just use sample data
            logger.warning(f"JSON parsing failed: {parse_error}")
            sample_data["llm_response"] = result[:200]

        logger.info(f"LLM extraction successful for document {doc_id}")

    except Exception as e:
        logger.warning(f"LLM extraction failed for document {doc_id}: {e}", exc_info=True)
        # Continue with sample data - add error details
        sample_data["llm_fehler"] = f"{type(e).__name__}: {str(e)}"

    # Store extracted data (either from LLM or sample)
    for field_name, field_value in sample_data.items():
        if field_value:
            data_entry = TaxCaseExtractedData(
                tax_case_id=case_id,
                document_id=doc_id,
                field_name=field_name,
                field_value=str(field_value),
                field_type="text",
                confidence=0.5,
                confirmed=False
            )
            db.add(data_entry)

    db.flush()
    logger.info(f"Stored {len(sample_data)} fields for document {doc_id}")


@router.get("/cases/{case_id}/documents", response_model=List[DocumentResponse])
async def list_documents(case_id: int, db: Session = Depends(get_db), user: User = Depends(get_demo_user)):
    """List all documents for a tax case"""
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Tax case not found")

    documents = db.query(TaxCaseDocument).filter(
        TaxCaseDocument.tax_case_id == case_id
    ).order_by(desc(TaxCaseDocument.created_at)).all()

    return [DocumentResponse(
        id=doc.id,
        tax_case_id=doc.tax_case_id,
        folder_id=doc.folder_id,
        filename=doc.filename,
        doc_type=doc.doc_type or "document",
        validated=doc.validated,
        created_at=doc.created_at
    ) for doc in documents]


@router.put("/cases/{case_id}/data")
async def update_extracted_data(
    case_id: int,
    data_update: ExtractedDataBulkUpdate,
    db: Session = Depends(get_db),
    user: User = Depends(get_demo_user)
):
    """Update extracted data for a tax case"""
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Tax case not found")

    # Update each field
    for field_name, field_value in data_update.data.items():
        existing = db.query(TaxCaseExtractedData).filter(
            TaxCaseExtractedData.tax_case_id == case_id,
            TaxCaseExtractedData.field_name == field_name
        ).first()

        if existing:
            if existing.field_value != str(field_value):
                existing.original_value = existing.field_value
                existing.field_value = str(field_value)
                existing.edited = True
        else:
            # Create new field
            new_data = TaxCaseExtractedData(
                tax_case_id=case_id,
                field_name=field_name,
                field_value=str(field_value),
                field_type="text",
                edited=True
            )
            db.add(new_data)

    db.commit()
    logger.info(f"Updated extracted data for case {case_id}")

    return {"success": True, "message": "Data updated successfully"}


@router.post("/cases/{case_id}/confirm")
async def confirm_case(case_id: int, db: Session = Depends(get_db), user: User = Depends(get_demo_user)):
    """Confirm and validate a tax case after review"""
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Tax case not found")

    # Mark case as validated
    case.validated = True
    case.status = "confirmed"

    # Mark all extracted data as confirmed
    db.query(TaxCaseExtractedData).filter(
        TaxCaseExtractedData.tax_case_id == case_id
    ).update({"confirmed": True})

    # Mark all documents as validated
    db.query(TaxCaseDocument).filter(
        TaxCaseDocument.tax_case_id == case_id
    ).update({"validated": True})

    db.commit()
    logger.info(f"Confirmed tax case {case_id}")

    return {"success": True, "message": "Case confirmed successfully"}


@router.get("/cases/{case_id}/export")
async def export_case_pdf(case_id: int, db: Session = Depends(get_db), user: User = Depends(get_demo_user)):
    """Export tax case data to PDF"""
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Tax case not found")

    # Get extracted data
    extracted_items = db.query(TaxCaseExtractedData).filter(
        TaxCaseExtractedData.tax_case_id == case_id
    ).all()

    # Create PDF in memory
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # Title
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 50, f"Steuerfall: {case.name}")

    # Metadata
    p.setFont("Helvetica", 10)
    p.drawString(50, height - 80, f"Status: {case.status}")
    p.drawString(50, height - 95, f"Erstellt: {case.created_at.strftime('%d.%m.%Y %H:%M')}")
    p.drawString(50, height - 110, f"Validiert: {'Ja' if case.validated else 'Nein'}")

    # Extracted Data
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, height - 140, "Extrahierte Daten:")

    p.setFont("Helvetica", 10)
    y_position = height - 160
    for item in extracted_items:
        if y_position < 50:
            p.showPage()
            y_position = height - 50

        p.drawString(50, y_position, f"{item.field_name}: {item.field_value}")
        y_position -= 15

    p.showPage()
    p.save()

    buffer.seek(0)
    logger.info(f"Exported PDF for case {case_id}")

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=fall_{case_id}_export.pdf"}
    )


@router.get("/documents/{doc_id}/download")
async def download_document(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_demo_user)):
    """Download a specific document"""
    doc = db.query(TaxCaseDocument).filter(TaxCaseDocument.id == doc_id).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Verify ownership
    case = db.query(TaxCase).filter(
        TaxCase.id == doc.tax_case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=403, detail="Access denied")

    # Check if file exists
    if not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="File not found on server")

    def file_iterator():
        with open(doc.file_path, "rb") as f:
            yield from f

    return StreamingResponse(
        file_iterator(),
        media_type="application/octet-stream",
        headers={"Content-Disposition": f"attachment; filename={doc.filename}"}
    )


# Health check
@router.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "ok", "service": "taxcases", "version": "1.0.0"}


# Debug endpoint
@router.get("/debug/check-services")
async def debug_check_services():
    """Debug: Check OCR and LLM services availability"""
    status = {
        "tesseract": False,
        "paddleocr": False,
        "ollama": False,
        "grok": False,
        "errors": {}
    }

    # Check Tesseract
    try:
        from PIL import Image
        import pytesseract
        status["tesseract"] = True
    except Exception as e:
        status["errors"]["tesseract"] = str(e)

    # Check PaddleOCR with detailed error tracking
    paddle_status = {"available": False, "import_error": None, "init_error": None}
    try:
        # Test 1: Can we import paddleocr?
        import paddleocr
        paddle_status["paddleocr_imported"] = True
        paddle_status["paddleocr_location"] = str(paddleocr.__file__)

        # Test 2: Can we import paddlepaddle?
        import paddle
        paddle_status["paddle_imported"] = True
        paddle_status["paddle_version"] = paddle.__version__

        # Test 3: Can we get singleton instance?
        from backend.services.application_service import get_paddle_ocr, PADDLE_AVAILABLE
        paddle_status["PADDLE_AVAILABLE"] = PADDLE_AVAILABLE

        if PADDLE_AVAILABLE:
            ocr = get_paddle_ocr()
            if ocr:
                status["paddleocr"] = True
                status["paddleocr_details"] = paddle_status
            else:
                status["paddleocr"] = False
                paddle_status["init_error"] = "get_paddle_ocr() returned None"
                status["errors"]["paddleocr"] = paddle_status
        else:
            status["paddleocr"] = False
            status["errors"]["paddleocr"] = paddle_status
    except ImportError as e:
        status["paddleocr"] = False
        paddle_status["import_error"] = f"ImportError: {str(e)}"
        status["errors"]["paddleocr"] = paddle_status
    except Exception as e:
        status["paddleocr"] = False
        paddle_status["other_error"] = f"{type(e).__name__}: {str(e)}"
        status["errors"]["paddleocr"] = paddle_status

    # Check Ollama
    try:
        from backend.services.llm_gateway import llm_gateway
        result = llm_gateway.generate(prompt="Test", provider="ollama", max_tokens=10, timeout=10)
        status["ollama"] = True
        status["ollama_response"] = result.get("response", "")[:50]
    except Exception as e:
        status["errors"]["ollama"] = str(e)

    # Check Grok
    try:
        from backend.services.llm_gateway import llm_gateway
        from backend.config import settings
        if settings.GROK_API_KEY and settings.GROK_API_KEY.strip():
            status["grok"] = "API key configured"
        else:
            status["grok"] = False
            status["errors"]["grok"] = "API key not configured"
    except Exception as e:
        status["errors"]["grok"] = str(e)

    return status


@router.get("/debug/pip-list")
async def debug_pip_list():
    """Debug: Check installed packages"""
    import subprocess
    try:
        result = subprocess.run(['pip', 'list'], capture_output=True, text=True, timeout=10)
        lines = result.stdout.split('\n')

        # Filter for relevant packages
        relevant = []
        for line in lines:
            if any(pkg in line.lower() for pkg in ['paddle', 'tesseract', 'pillow', 'numpy', 'opencv']):
                relevant.append(line.strip())

        return {
            "success": True,
            "relevant_packages": relevant,
            "full_output_lines": len(lines)
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@router.post("/debug/test-llm")
async def debug_test_llm(data: dict):
    """Debug: Test LLM extraction with simple prompt"""
    try:
        from backend.services.llm_gateway import llm_gateway

        provider = data.get("provider", "grok")
        prompt = data.get("prompt", "Extract the invoice number from this text: Invoice RE-2026-001")

        result = llm_gateway.generate(
            prompt=prompt,
            provider=provider,
            max_tokens=200,
            timeout=60
        )

        return {
            "success": True,
            "provider": provider,
            "response": result.get("response", ""),
            "model": result.get("model", ""),
            "usage": result.get("usage", {})
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }


# FREE TAB ENDPOINTS

@router.post("/free/upload-and-preview")
async def free_upload_and_preview(
    email: str = Form(...),
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_demo_user)
):
    """Upload image, apply rotation correction, return preview"""
    import base64
    from datetime import datetime

    # Create temporary case
    case_name = f"Free Upload Preview - {email} - {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    new_case = TaxCase(
        user_id=user.id,
        name=case_name,
        status="preview",
        notes="Bildvorschau - noch nicht verarbeitet"
    )
    db.add(new_case)
    db.flush()

    case_id = new_case.id
    upload_dir = f"/tmp/tax_uploads/{case_id}"
    os.makedirs(upload_dir, exist_ok=True)

    processed_images = []

    try:
        for file in files:
            # Save original file
            original_path = os.path.join(upload_dir, f"original_{file.filename}")
            with open(original_path, "wb") as f:
                content = await file.read()
                f.write(content)

            # Apply rotation correction if it's an image
            if file.filename.lower().endswith(('.jpg', '.jpeg', '.png', '.tiff', '.bmp')):
                try:
                    from backend.services.application_service import detect_and_correct_rotation, CV2_AVAILABLE
                    import cv2

                    # Convert original to base64
                    with open(original_path, "rb") as f:
                        original_base64 = base64.b64encode(f.read()).decode('utf-8')

                    if CV2_AVAILABLE:
                        # Apply rotation correction and get OCR qualities
                        # Returns COLOR image (BGR format) + original quality + processed quality
                        corrected_array, original_quality, processed_quality = detect_and_correct_rotation(original_path)

                        # Save corrected image
                        corrected_path = os.path.join(upload_dir, file.filename)
                        cv2.imwrite(corrected_path, corrected_array)

                        # Convert corrected to base64
                        _, buffer = cv2.imencode('.jpg', corrected_array)
                        processed_base64 = base64.b64encode(buffer).decode('utf-8')

                        # Check if quality is good enough (>= 90%)
                        quality_ok = processed_quality >= 90.0
                        quality_message = None
                        if not quality_ok:
                            quality_message = f"OCR-Qualität zu niedrig ({processed_quality:.1f}%). Bitte machen Sie ein normales, scharfes Foto der Rechnung."

                        # Extract quick OCR preview (first 200 chars) for debugging
                        ocr_preview = ""
                        try:
                            from PIL import Image as PILImage
                            import pytesseract
                            # Convert BGR to RGB for PIL
                            corrected_rgb_pil = cv2.cvtColor(corrected_array, cv2.COLOR_BGR2RGB)
                            pil_img = PILImage.fromarray(corrected_rgb_pil)
                            european_langs = 'deu+eng+spa+fra+ita'
                            preview_text = pytesseract.image_to_string(pil_img, lang=european_langs, config='--psm 1')
                            ocr_preview = preview_text[:200].strip() if preview_text else ""
                        except Exception as ocr_err:
                            logger.warning(f"OCR preview failed: {ocr_err}")

                        processed_images.append({
                            "filename": file.filename,
                            "original_preview": f"data:image/jpeg;base64,{original_base64}",
                            "processed_preview": f"data:image/jpeg;base64,{processed_base64}",
                            "path": corrected_path,
                            "original_ocr_quality": round(original_quality, 1),
                            "processed_ocr_quality": round(processed_quality, 1),
                            "ocr_quality": round(processed_quality, 1),  # Keep for backward compatibility
                            "quality_ok": quality_ok,
                            "quality_message": quality_message,
                            "ocr_preview": ocr_preview
                        })
                    else:
                        # No OpenCV, just save original
                        logger.warning("OpenCV not available, skipping rotation correction")
                        processed_images.append({
                            "filename": file.filename,
                            "original_preview": f"data:image/jpeg;base64,{original_base64}",
                            "processed_preview": f"data:image/jpeg;base64,{original_base64}",
                            "path": original_path,
                            "ocr_quality": 0.0,
                            "quality_ok": False,
                            "quality_message": "Rotation correction not available"
                        })
                except Exception as img_err:
                    logger.error(f"Image processing failed: {img_err}")
                    # Fallback: Return original only
                    with open(original_path, "rb") as f:
                        original_base64 = base64.b64encode(f.read()).decode('utf-8')
                    processed_images.append({
                        "filename": file.filename,
                        "original_preview": f"data:image/jpeg;base64,{original_base64}",
                        "processed_preview": f"data:image/jpeg;base64,{original_base64}",
                        "path": original_path,
                        "ocr_quality": 0.0,
                        "quality_ok": False,
                        "quality_message": f"Bildverarbeitung fehlgeschlagen: {str(img_err)}",
                        "error": str(img_err)
                    })
            elif file.filename.lower().endswith('.pdf'):
                # PDF file - try direct text extraction first (no OCR needed)
                try:
                    from PyPDF2 import PdfReader

                    logger.info(f"Attempting direct PDF text extraction: {file.filename}")

                    # Try to extract text directly from PDF
                    pdf_reader = PdfReader(original_path)
                    extracted_text = ""

                    for page_num, page in enumerate(pdf_reader.pages, start=1):
                        page_text = page.extract_text()
                        if page_text:
                            extracted_text += f"\n--- Seite {page_num} ---\n{page_text}"

                    # Check if we got meaningful text (more than 50 characters)
                    if len(extracted_text.strip()) > 50:
                        logger.info(f"✅ PDF text extracted directly ({len(extracted_text)} chars, no OCR needed)")

                        word_count = len(extracted_text.split())
                        processed_images.append({
                            "filename": file.filename,
                            "message": f"PDF erfolgreich verarbeitet ({word_count} Wörter, {len(pdf_reader.pages)} Seite(n))",
                            "quality_ok": True,
                            "ocr_quality": 100.0,  # Perfect quality - no OCR needed
                            "document_type": "pdf_text",
                            "text_length": len(extracted_text),
                            "page_count": len(pdf_reader.pages),
                            "method": "direct_extraction"
                        })

                        # Store text for later processing
                        extracted_texts[file.filename] = extracted_text
                        continue  # Skip OCR processing
                    else:
                        logger.info(f"⚠️ PDF has no extractable text, falling back to OCR")

                except Exception as extract_err:
                    logger.warning(f"Direct PDF extraction failed: {str(extract_err)}, falling back to OCR")

                # Fallback: PDF is scanned/image-based - use OCR
                try:
                    from pdf2image import convert_from_path
                    from backend.services.application_service import detect_and_correct_rotation, CV2_AVAILABLE
                    import cv2

                    logger.info(f"Converting PDF to images for OCR: {file.filename}")

                    # Convert PDF pages to images (300 DPI for good quality)
                    pdf_images = convert_from_path(original_path, dpi=300)
                    logger.info(f"PDF has {len(pdf_images)} page(s)")

                    # Process each page
                    for page_num, pdf_page_image in enumerate(pdf_images, start=1):
                        try:
                            # Save PDF page as temporary image
                            page_filename = f"{os.path.splitext(file.filename)[0]}_Seite_{page_num}.jpg"
                            page_path = os.path.join(upload_dir, f"pdf_page_{page_num}_{file.filename}.jpg")
                            pdf_page_image.save(page_path, 'JPEG', quality=95)

                            # Convert to base64 for preview
                            with open(page_path, "rb") as f:
                                original_base64 = base64.b64encode(f.read()).decode('utf-8')

                            if CV2_AVAILABLE:
                                # Apply rotation correction and get OCR qualities
                                corrected_array, original_quality, processed_quality = detect_and_correct_rotation(page_path)

                                # Save corrected image
                                corrected_path = os.path.join(upload_dir, page_filename)
                                cv2.imwrite(corrected_path, corrected_array)

                                # Convert corrected to base64
                                _, buffer = cv2.imencode('.jpg', corrected_array)
                                processed_base64 = base64.b64encode(buffer).decode('utf-8')

                                # Check if quality is good enough (>= 90%)
                                quality_ok = processed_quality >= 90.0
                                quality_message = None
                                if not quality_ok:
                                    quality_message = f"OCR-Qualität zu niedrig ({processed_quality:.1f}%). Bitte verwenden Sie ein besseres PDF."

                                # Extract quick OCR preview
                                ocr_preview = ""
                                try:
                                    from PIL import Image as PILImage
                                    import pytesseract
                                    corrected_rgb_pil = cv2.cvtColor(corrected_array, cv2.COLOR_BGR2RGB)
                                    pil_img = PILImage.fromarray(corrected_rgb_pil)
                                    european_langs = 'deu+eng+spa+fra+ita'
                                    preview_text = pytesseract.image_to_string(pil_img, lang=european_langs, config='--psm 1')
                                    ocr_preview = preview_text[:200].strip() if preview_text else ""
                                except Exception as ocr_err:
                                    logger.warning(f"OCR preview failed for PDF page {page_num}: {ocr_err}")

                                processed_images.append({
                                    "filename": page_filename,
                                    "original_preview": f"data:image/jpeg;base64,{original_base64}",
                                    "processed_preview": f"data:image/jpeg;base64,{processed_base64}",
                                    "path": corrected_path,
                                    "original_ocr_quality": round(original_quality, 1),
                                    "processed_ocr_quality": round(processed_quality, 1),
                                    "ocr_quality": round(processed_quality, 1),
                                    "quality_ok": quality_ok,
                                    "quality_message": quality_message,
                                    "ocr_preview": ocr_preview,
                                    "source_pdf": file.filename,
                                    "page_number": page_num
                                })
                            else:
                                # No OpenCV, just save original
                                logger.warning("OpenCV not available, skipping rotation correction for PDF page")
                                processed_images.append({
                                    "filename": page_filename,
                                    "original_preview": f"data:image/jpeg;base64,{original_base64}",
                                    "processed_preview": f"data:image/jpeg;base64,{original_base64}",
                                    "path": page_path,
                                    "ocr_quality": 0.0,
                                    "quality_ok": False,
                                    "quality_message": "Rotation correction not available",
                                    "source_pdf": file.filename,
                                    "page_number": page_num
                                })
                        except Exception as page_err:
                            logger.error(f"Failed to process PDF page {page_num}: {page_err}")
                            # Add error entry for this page
                            processed_images.append({
                                "filename": f"{os.path.splitext(file.filename)[0]}_Seite_{page_num}.jpg",
                                "ocr_quality": 0.0,
                                "quality_ok": False,
                                "quality_message": f"Fehler bei Seite {page_num}: {str(page_err)}",
                                "error": str(page_err),
                                "source_pdf": file.filename,
                                "page_number": page_num
                            })

                except Exception as pdf_err:
                    logger.error(f"PDF conversion failed: {pdf_err}")
                    processed_images.append({
                        "filename": file.filename,
                        "message": f"PDF-Konvertierung fehlgeschlagen: {str(pdf_err)}",
                        "error": str(pdf_err)
                    })
            elif file.filename.lower().endswith(('.odt', '.docx', '.doc', '.txt')):
                # Text documents - direct text extraction without OCR
                try:
                    file_path = os.path.join(upload_dir, file.filename)
                    with open(file_path, "wb") as f:
                        f.write(content)

                    # Extract text based on file type
                    if file.filename.lower().endswith('.txt'):
                        # TXT - simple file read
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            extracted_text = f.read()
                    elif file.filename.lower().endswith(('.docx', '.doc')):
                        # DOCX/DOC - use python-docx (synchronous, no asyncio issues)
                        # python-docx can sometimes read .doc files too
                        try:
                            from docx import Document
                            doc = Document(file_path)
                            extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            logger.info(f"Extracted {len(extracted_text)} chars from {file.filename} using python-docx")
                        except Exception as docx_err:
                            logger.error(f"Document extraction failed for {file.filename}: {docx_err}")

                            # Fallback: Convert to PDF then to images and OCR (especially for old .doc format)
                            if file.filename.lower().endswith('.doc'):
                                logger.info(f"Trying LibreOffice conversion for {file.filename}")
                                try:
                                    # Convert DOC to PDF using LibreOffice
                                    import subprocess
                                    pdf_path = file_path.replace('.doc', '.pdf')
                                    result = subprocess.run(
                                        ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', upload_dir, file_path],
                                        capture_output=True,
                                        timeout=30
                                    )

                                    if os.path.exists(pdf_path):
                                        logger.info(f"Converted {file.filename} to PDF, now converting to images...")
                                        # Convert PDF to images
                                        from pdf2image import convert_from_path
                                        images = convert_from_path(pdf_path, dpi=300)

                                        # OCR on each page
                                        extracted_texts = []
                                        for idx, img in enumerate(images):
                                            img_path = os.path.join(upload_dir, f"{file.filename}_page_{idx+1}.png")
                                            img.save(img_path, 'PNG')

                                            # OCR
                                            import pytesseract
                                            text = pytesseract.image_to_string(img, lang='deu+eng')
                                            extracted_texts.append(text)
                                            logger.info(f"OCR extracted {len(text)} chars from page {idx+1}")

                                        extracted_text = "\n\n--- PAGE BREAK ---\n\n".join(extracted_texts)
                                        logger.info(f"Total OCR extraction: {len(extracted_text)} chars from {len(images)} pages")
                                    else:
                                        raise Exception("LibreOffice conversion failed")

                                except Exception as convert_err:
                                    logger.error(f"Conversion and OCR failed for {file.filename}: {convert_err}")
                                    extracted_text = f"[Error: Could not extract from .doc file. Please convert to .docx, PDF, or TXT.]"
                            else:
                                extracted_text = f"[Error extracting document. File may be corrupted or in unsupported format. Try converting to PDF or TXT.]"
                    elif file.filename.lower().endswith('.odt'):
                        # ODT - try odfpy
                        try:
                            from odf import text, teletype
                            from odf.opendocument import load
                            textdoc = load(file_path)
                            allparas = textdoc.getElementsByType(text.P)
                            extracted_text = "\n".join([teletype.extractText(p) for p in allparas])
                            logger.info(f"Extracted {len(extracted_text)} chars from ODT")
                        except Exception as odt_err:
                            logger.error(f"ODT extraction failed: {odt_err}")
                            extracted_text = f"[Error extracting ODT: {str(odt_err)}]"
                    else:
                        # Unsupported format
                        extracted_text = "[Error: Unsupported document format]"

                    # Check if extraction was successful
                    if extracted_text and not extracted_text.startswith('[Error'):
                        word_count = len(extracted_text.split())
                        processed_images.append({
                            "filename": file.filename,
                            "message": f"Textdokument erfolgreich verarbeitet ({word_count} Wörter extrahiert)",
                            "path": file_path,
                            "text_preview": extracted_text[:300],  # First 300 chars
                            "word_count": word_count,
                            "quality_ok": True,
                            "ocr_quality": 100.0,  # Text documents don't need OCR
                            "document_type": "text"
                        })
                    else:
                        processed_images.append({
                            "filename": file.filename,
                            "message": f"Textextraktion fehlgeschlagen: {extracted_text}",
                            "quality_ok": False,
                            "error": extracted_text
                        })
                except Exception as doc_err:
                    logger.error(f"Document processing failed for {file.filename}: {doc_err}")
                    processed_images.append({
                        "filename": file.filename,
                        "message": f"Dokumentverarbeitung fehlgeschlagen: {str(doc_err)}",
                        "error": str(doc_err),
                        "quality_ok": False
                    })
            else:
                # Unsupported file type
                processed_images.append({
                    "filename": file.filename,
                    "message": "Nicht unterstütztes Dateiformat. Unterstützt: PDF, JPG, PNG, ODT, DOCX, DOC, TXT",
                    "quality_ok": False
                })

        db.commit()

        return {
            "success": True,
            "case_id": case_id,
            "processed_images": processed_images,
            "message": f"{len(processed_images)} Bilder verarbeitet und rotationskorrigiert"
        }

    except Exception as e:
        db.rollback()
        logger.error(f"Upload preview failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


@router.post("/free/extract")
async def free_extract(
    case_id: int = Form(...),
    use_local_llm: str = Form("false"),  # Ignored - always use Grok for cost efficiency
    db: Session = Depends(get_db),
    user: User = Depends(get_demo_user)
):
    """Extract data from previously uploaded and processed images - always uses Grok 4.1 Fast"""
    # Always use Grok for MVP Tax Spain (cost optimization with grok-4-1-fast)
    prefer_local = False
    llm_mode = "Grok 4.1 Fast API"

    # Get case
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    # Get uploaded files
    upload_dir = f"/tmp/tax_uploads/{case_id}"
    if not os.path.exists(upload_dir):
        raise HTTPException(status_code=404, detail="Upload directory not found")

    # Find processed images
    files = [f for f in os.listdir(upload_dir) if not f.startswith("original_")]
    if not files:
        raise HTTPException(status_code=404, detail="No processed files found")

    extracted_data = {
        "verarbeitungsmodus": llm_mode,
        "ocr_engine": "Tesseract OCR",
        "email": case.name.split(" - ")[1] if " - " in case.name else "unknown",
        "dateiname": ", ".join(files)
    }

    try:
        # Extract text from all files (images, PDFs, text documents)
        all_content = []

        for filename in files:
            file_path = os.path.join(upload_dir, filename)

            # Check file type
            is_docx = filename.lower().endswith('.docx')
            is_odt = filename.lower().endswith('.odt')
            is_txt = filename.lower().endswith('.txt')
            is_doc = filename.lower().endswith('.doc')  # Legacy Word format
            is_pdf = filename.lower().endswith('.pdf')

            if DocumentParser:
                try:
                    parser = DocumentParser()

                    if is_txt:
                        # TXT - simple file read
                        logger.info(f"Reading TXT file: {filename}")
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            doc_content = f.read()
                        logger.info(f"Read {len(doc_content)} characters from TXT {filename}")
                    elif is_docx or is_doc:
                        # DOCX/DOC - use python-docx (can handle both)
                        logger.info(f"Extracting Word document: {filename}")
                        try:
                            from docx import Document
                            doc = Document(file_path)
                            doc_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            logger.info(f"Extracted {len(doc_content)} characters from {filename} using python-docx")
                        except Exception as docx_err:
                            logger.error(f"Word document extraction failed: {docx_err}")

                            # Fallback: Convert to PDF then to images and OCR (especially for old .doc format)
                            if is_doc:
                                logger.info(f"Trying LibreOffice conversion for {filename}")
                                try:
                                    # Convert DOC to PDF using LibreOffice
                                    import subprocess
                                    pdf_path = file_path.replace('.doc', '.pdf')
                                    result = subprocess.run(
                                        ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', upload_dir, file_path],
                                        capture_output=True,
                                        timeout=30
                                    )

                                    if os.path.exists(pdf_path):
                                        logger.info(f"Converted {filename} to PDF, now converting to images...")
                                        # Convert PDF to images
                                        from pdf2image import convert_from_path
                                        images = convert_from_path(pdf_path, dpi=300)

                                        # OCR on each page
                                        extracted_texts = []
                                        for idx, img in enumerate(images):
                                            img_path = os.path.join(upload_dir, f"{filename}_page_{idx+1}.png")
                                            img.save(img_path, 'PNG')

                                            # OCR
                                            import pytesseract
                                            text = pytesseract.image_to_string(img, lang='deu+eng')
                                            extracted_texts.append(text)
                                            logger.info(f"OCR extracted {len(text)} chars from page {idx+1}")

                                        doc_content = "\n\n--- PAGE BREAK ---\n\n".join(extracted_texts)
                                        logger.info(f"Total OCR extraction: {len(doc_content)} chars from {len(images)} pages")
                                    else:
                                        raise Exception("LibreOffice conversion failed")

                                except Exception as convert_err:
                                    logger.error(f"Conversion and OCR failed for {filename}: {convert_err}")
                                    doc_content = f"[Error: Could not extract from .doc file. Please convert to .docx, PDF, or TXT.]"
                            else:
                                doc_content = f"[Error extracting document. File may be corrupted or in unsupported format. Try converting to PDF or TXT.]"
                    elif is_odt:
                        # ODT - use odfpy
                        logger.info(f"Extracting ODT: {filename}")
                        try:
                            from odf import text, teletype
                            from odf.opendocument import load
                            textdoc = load(file_path)
                            allparas = textdoc.getElementsByType(text.P)
                            doc_content = "\n".join([teletype.extractText(p) for p in allparas])
                            logger.info(f"Extracted {len(doc_content)} characters from ODT")
                        except Exception as odt_err:
                            logger.error(f"ODT extraction failed: {odt_err}")
                            doc_content = f"[Error extracting ODT: {str(odt_err)}]"
                    elif is_pdf:
                        # PDF - try direct text extraction first
                        logger.info(f"Extracting text from PDF: {filename}")
                        try:
                            from PyPDF2 import PdfReader
                            pdf_reader = PdfReader(file_path)
                            extracted_text = ""
                            for page in pdf_reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    extracted_text += page_text + "\n"

                            if len(extracted_text.strip()) > 50:
                                doc_content = extracted_text
                                logger.info(f"Extracted {len(doc_content)} characters from PDF {filename} (direct)")
                            else:
                                # Fallback to OCR
                                doc_content = parser._parse_image_tesseract(
                                    open(file_path, 'rb').read(),
                                    file_path=None
                                )
                                logger.info(f"Extracted {len(doc_content)} characters from PDF {filename} (OCR)")
                        except Exception as pdf_err:
                            logger.warning(f"PDF direct extraction failed, using OCR: {pdf_err}")
                            doc_content = parser._parse_image_tesseract(
                                open(file_path, 'rb').read(),
                                file_path=None
                            )
                    else:
                        # Image - use Tesseract (already rotated/corrected)
                        doc_content = parser._parse_image_tesseract(
                            open(file_path, 'rb').read(),
                            file_path=None  # Skip rotation correction
                        )
                        logger.info(f"Extracted {len(doc_content)} characters from image {filename}")

                except Exception as parse_error:
                    logger.warning(f"Extraction failed for {filename}: {parse_error}")
                    doc_content = f"Dokument: {filename} (Extraktion fehlgeschlagen: {str(parse_error)})"
            else:
                doc_content = f"Dokument: {filename}"

            all_content.append(f"--- {filename} ---\n{doc_content}\n")

        combined_content = "\n\n".join(all_content)

        # Initialize result variable for debug_info
        result = ""

        # LLM Extraction (same as original /free/upload)
        try:
            from backend.services.llm_gateway import llm_gateway

            prompt = f"""Extrahiere folgende Informationen aus diesem Dokument für Import auf die Kanarischen Inseln (H7-Formular):

1. ALLGEMEINE SENDUNGSDATEN:
- art_der_sendung: B2C (Business to Consumer) oder C2C (Consumer to Consumer)
- warenwert_gesamt_eur: Gesamtwert der Waren in EUR
- waehrung: Währung (meist EUR)
- versandkosten: Versandkosten
- versicherungskosten: Versicherungskosten (optional)
- gesamtbetrag_fuer_zoll: Summe aus Wert + Versand + Versicherung
- art_der_lieferung: Kauf oder Geschenk

2. ABSENDER (Versender):
- absender_name: Name oder Firma des Absenders
- absender_strasse: Straße und Hausnummer
- absender_plz: Postleitzahl
- absender_ort: Ort
- absender_land: Land (ISO-Code)
- absender_email: E-Mail (optional)
- absender_telefon: Telefonnummer (optional)

3. EMPFÄNGER (Kanarische Inseln):
- empfaenger_name: Name des Empfängers
- empfaenger_strasse: Straße und Hausnummer
- empfaenger_plz: Postleitzahl
- empfaenger_ort: Ort
- empfaenger_insel: Welche kanarische Insel (wichtig!)
- empfaenger_nif_nie_cif: NIF/NIE/CIF Nummer (sehr wichtig!)
- empfaenger_email: E-Mail
- empfaenger_telefon: Telefonnummer

4. WARENPOSITIONEN (mindestens 1):
- position_1_beschreibung: Klare Warenbeschreibung
- position_1_anzahl: Stückzahl
- position_1_stueckpreis: Preis pro Stück
- position_1_gesamtwert: Gesamtwert der Position
- position_1_ursprungsland: Ursprungsland (ISO-Code)
- position_1_zolltarifnummer: 6-stellige Zolltarifnummer (optional)
- position_1_gewicht: Gewicht (optional)
- position_1_zustand: Neu oder gebraucht (optional)

5. RECHNUNGS-/WERTNACHWEIS:
- rechnungsnummer: Nummer der Rechnung
- rechnungsdatum: Datum der Rechnung
- mehrwertsteuer_ausgewiesen: Ja/Nein, wurde MwSt ausgewiesen?

6. ZUSÄTZLICHE ANGABEN:
- zahlungsnachweis: Art des Zahlungsnachweises (PayPal, Karte, etc.)
- bemerkungen: Weitere Anmerkungen

Dokumenteninhalt:
{combined_content[:8000]}

Gib die extrahierten Daten als JSON-Objekt zurück.
Verwende exakt die Feldnamen wie oben angegeben.
Falls ein Feld nicht gefunden wird, lasse es weg oder gib null zurück.
Bei mehreren Warenpositionen verwende position_2_*, position_3_* usw.
"""

            provider = "ollama" if prefer_local else "grok"
            timeout = 240 if provider == "ollama" else 60

            logger.info(f"Sending to LLM: {len(combined_content)} chars of text")
            logger.info(f"Text preview: {combined_content[:500]}")

            result_dict = llm_gateway.generate(prompt=prompt, provider=provider, max_tokens=4000, timeout=timeout)
            result = result_dict.get("response", "")

            logger.info(f"LLM extraction completed with {llm_mode}")
            logger.info(f"LLM response length: {len(result)} chars")
            logger.info(f"LLM response preview: {result[:500]}")

            # Parse JSON
            try:
                llm_data = llm_gateway.parse_json_response(result)
                # Flatten nested objects
                flattened = {}
                for key, value in llm_data.items():
                    if isinstance(value, dict):
                        for nested_key, nested_value in value.items():
                            flattened[nested_key] = nested_value
                    else:
                        flattened[key] = value

                # Count non-null fields
                non_null_count = sum(1 for v in flattened.values() if v not in (None, '', 'null'))
                extracted_data.update(flattened)
                logger.info(f"Successfully parsed {len(flattened)} fields ({non_null_count} non-null)")
                logger.info(f"Sample fields: {dict(list(flattened.items())[:5])}")
            except (json.JSONDecodeError, ValueError) as parse_error:
                logger.warning(f"JSON parsing failed: {parse_error}")
                extracted_data["llm_antwort"] = result[:1000]
                extracted_data["inhalt_vorschau"] = combined_content[:500]

        except Exception as llm_error:
            logger.error(f"LLM extraction failed: {llm_error}", exc_info=True)
            extracted_data["fehler"] = f"LLM-Extraktion fehlgeschlagen: {type(llm_error).__name__}"
            extracted_data["fehler_details"] = str(llm_error)
            extracted_data["dokument_vorschau"] = combined_content[:800]

        # Save to database
        for field_name, field_value in extracted_data.items():
            if field_value:
                data_entry = TaxCaseExtractedData(
                    tax_case_id=case_id,
                    field_name=field_name,
                    field_value=str(field_value),
                    field_type="text",
                    confidence=0.7,
                    confirmed=False
                )
                db.add(data_entry)

        case.status = "validated"
        db.commit()

        return {
            "success": True,
            "extracted_data": extracted_data,
            "message": f"Daten erfolgreich extrahiert ({llm_mode})",
            "debug_info": {
                "extracted_text": combined_content[:5000],  # First 5000 chars of extracted text
                "llm_response": result[:5000] if 'result' in locals() else None  # First 5000 chars of LLM response
            }
        }

    except Exception as e:
        db.rollback()
        logger.error(f"Extraction failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")


@router.post("/free/upload")
async def free_upload(
    email: str = Form(...),
    files: List[UploadFile] = File(...),
    use_local_llm: str = Form("true"),
    ocr_engine: str = Form("tesseract"),
    db: Session = Depends(get_db),
    user: User = Depends(get_demo_user)
):
    """Free tier: Upload documents and extract data - SAVES to database"""
    prefer_local = use_local_llm.lower() == "true"
    llm_mode = "Lokales LLM (DSGVO)" if prefer_local else "Grok API"
    ocr_mode = "PaddleOCR" if ocr_engine.lower() == "paddle" else "Tesseract OCR"

    # Create a case for this upload (so it appears in Premium tab)
    from datetime import datetime
    case_name = f"Free Upload - {email} - {datetime.now().strftime('%d.%m.%Y %H:%M')}"

    new_case = TaxCase(
        user_id=user.id,
        name=case_name,
        status="processing",
        notes=f"Erstellt über Free Tab mit {llm_mode} und {ocr_mode}"
    )
    db.add(new_case)
    db.flush()

    case_id = new_case.id
    logger.info(f"Created case {case_id} for free upload from {email} using {ocr_mode}")

    # Create upload directory
    upload_dir = f"/tmp/tax_uploads/{case_id}"
    os.makedirs(upload_dir, exist_ok=True)

    extracted_data = {
        "verarbeitungsmodus": llm_mode,
        "ocr_engine": ocr_mode,
        "email": email
    }

    try:
        all_content = []
        filenames = []

        for file in files:
            filenames.append(file.filename)
            # Save file
            file_path = os.path.join(upload_dir, file.filename)
            with open(file_path, "wb") as f:
                content = await file.read()
                f.write(content)

            # Extract text with OCR/DocumentParser
            doc_content = ""
            if DocumentParser:
                try:
                    parser = DocumentParser()
                    doc_content = parser.parse(file_path, ocr_engine=ocr_engine.lower())
                    logger.info(f"{ocr_mode} extracted {len(doc_content)} characters from {file.filename}")
                except Exception as parse_error:
                    logger.warning(f"{ocr_mode} failed for {file.filename}: {parse_error}")
                    doc_content = f"Dokument: {file.filename} (OCR nicht verfügbar: {str(parse_error)})"
            else:
                doc_content = f"Dokument: {file.filename}"

            # Save document to database
            doc = TaxCaseDocument(
                tax_case_id=case_id,
                filename=file.filename,
                file_path=file_path,
                doc_type="document",
                content=doc_content,
                validated=False
            )
            db.add(doc)
            db.flush()

            all_content.append(f"--- {file.filename} ---\n{doc_content}\n")

        # Combine all document content
        combined_content = "\n\n".join(all_content)

        # Add filenames to extracted data
        extracted_data["dateiname"] = ", ".join(filenames) if len(filenames) > 0 else "Keine Datei"

        # Extract with LLM
        try:
            from backend.services.llm_gateway import llm_gateway

            prompt = f"""Extrahiere folgende Informationen aus diesem Dokument für Import auf die Kanarischen Inseln (H7-Formular):

1. ALLGEMEINE SENDUNGSDATEN:
- art_der_sendung: B2C (Business to Consumer) oder C2C (Consumer to Consumer)
- warenwert_gesamt_eur: Gesamtwert der Waren in EUR
- waehrung: Währung (meist EUR)
- versandkosten: Versandkosten
- versicherungskosten: Versicherungskosten (optional)
- gesamtbetrag_fuer_zoll: Summe aus Wert + Versand + Versicherung
- art_der_lieferung: Kauf oder Geschenk

2. ABSENDER (Versender):
- absender_name: Name oder Firma des Absenders
- absender_strasse: Straße und Hausnummer
- absender_plz: Postleitzahl
- absender_ort: Ort
- absender_land: Land (ISO-Code)
- absender_email: E-Mail (optional)
- absender_telefon: Telefonnummer (optional)

3. EMPFÄNGER (Kanarische Inseln):
- empfaenger_name: Name des Empfängers
- empfaenger_strasse: Straße und Hausnummer
- empfaenger_plz: Postleitzahl
- empfaenger_ort: Ort
- empfaenger_insel: Welche kanarische Insel (wichtig!)
- empfaenger_nif_nie_cif: NIF/NIE/CIF Nummer (sehr wichtig!)
- empfaenger_email: E-Mail
- empfaenger_telefon: Telefonnummer

4. WARENPOSITIONEN (mindestens 1):
- position_1_beschreibung: Klare Warenbeschreibung
- position_1_anzahl: Stückzahl
- position_1_stueckpreis: Preis pro Stück
- position_1_gesamtwert: Gesamtwert der Position
- position_1_ursprungsland: Ursprungsland (ISO-Code)
- position_1_zolltarifnummer: 6-stellige Zolltarifnummer (optional)
- position_1_gewicht: Gewicht (optional)
- position_1_zustand: Neu oder gebraucht (optional)

5. RECHNUNGS-/WERTNACHWEIS:
- rechnungsnummer: Nummer der Rechnung
- rechnungsdatum: Datum der Rechnung
- mehrwertsteuer_ausgewiesen: Ja/Nein, wurde MwSt ausgewiesen?

6. ZUSÄTZLICHE ANGABEN:
- zahlungsnachweis: Art des Zahlungsnachweises (PayPal, Karte, etc.)
- bemerkungen: Weitere Anmerkungen

Dokumenteninhalt:
{combined_content[:8000]}

Gib die extrahierten Daten als JSON-Objekt zurück.
Verwende exakt die Feldnamen wie oben angegeben.
Falls ein Feld nicht gefunden wird, lasse es weg oder gib null zurück.
Bei mehreren Warenpositionen verwende position_2_*, position_3_* usw.
"""

            # Convert prefer_local to provider
            provider = "ollama" if prefer_local else "grok"
            # Ollama on Railway is CPU-only and very slow, needs long timeout
            timeout = 240 if provider == "ollama" else 60  # 240 seconds for Ollama, 60 seconds for Grok
            result_dict = llm_gateway.generate(prompt=prompt, provider=provider, max_tokens=4000, timeout=timeout)
            result = result_dict.get("response", "")

            logger.info(f"LLM extraction completed with {llm_mode} using {provider}")

            # Parse LLM response
            try:
                # Use LLM gateway's robust JSON parser (handles markdown code blocks)
                llm_data = llm_gateway.parse_json_response(result)
                # Flatten nested objects (e.g., ABSENDER.absender_name -> absender_name)
                flattened = {}
                for key, value in llm_data.items():
                    if isinstance(value, dict):
                        # Flatten nested dict
                        for nested_key, nested_value in value.items():
                            flattened[nested_key] = nested_value
                    else:
                        flattened[key] = value
                extracted_data.update(flattened)
                logger.info(f"Successfully parsed {len(flattened)} fields from LLM response")
            except (json.JSONDecodeError, ValueError) as parse_error:
                # If parsing fails, use fallback
                logger.warning(f"LLM response parsing failed: {parse_error}")
                extracted_data["llm_antwort"] = result[:1000]
                # Add some basic extracted info
                extracted_data["inhalt_vorschau"] = combined_content[:500]

        except Exception as llm_error:
            logger.error(f"LLM extraction failed: {llm_error}", exc_info=True)
            # Fallback: provide document preview with detailed error
            extracted_data["fehler"] = f"LLM-Extraktion fehlgeschlagen: {type(llm_error).__name__}"
            extracted_data["fehler_details"] = str(llm_error)
            extracted_data["dokument_vorschau"] = combined_content[:800]
            extracted_data["hinweis"] = "Bitte überprüfen und manuell ausfüllen"

        # Ensure we have some data
        if len(extracted_data) <= 2:  # Only mode and email
            extracted_data["status"] = "Dokumente hochgeladen"
            extracted_data["dokument_inhalt"] = combined_content[:500]
            extracted_data["hinweis"] = "Bitte Daten manuell überprüfen und ergänzen"

        # Save extracted data to database
        for field_name, field_value in extracted_data.items():
            if field_value:
                data_entry = TaxCaseExtractedData(
                    tax_case_id=case_id,
                    field_name=field_name,
                    field_value=str(field_value),
                    field_type="text",
                    confidence=0.7,
                    confirmed=False
                )
                db.add(data_entry)

        # Mark case as validated
        new_case.status = "validated"
        db.commit()

        logger.info(f"Saved case {case_id} with {len(extracted_data)} extracted fields to database")

        return {
            "success": True,
            "extracted_data": extracted_data,
            "case_id": case_id,
            "message": f"Dokumente erfolgreich verarbeitet und gespeichert ({llm_mode})"
        }

    except Exception as e:
        logger.error(f"Error in free upload: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/free/create-document")
async def free_create_document(
    data: dict,
    db: Session = Depends(get_db)
):
    """Free tier: Create PDF document from extracted data"""
    try:
        email = data.get("email", "unknown@email.com")
        extracted_data = data.get("data", {})

        # Create PDF
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        # Title
        p.setFont("Helvetica-Bold", 20)
        p.drawString(50, height - 50, "Steuerdokument - Tax Document")

        # Email and date
        p.setFont("Helvetica", 10)
        p.drawString(50, height - 80, f"Erstellt für / Created for: {email}")
        p.drawString(50, height - 95, f"Datum / Date: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

        # Separator line
        p.line(50, height - 105, width - 50, height - 105)

        # Data section
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, height - 130, "Extrahierte Daten / Extracted Data:")

        p.setFont("Helvetica", 10)
        y_position = height - 155

        # Sort data to show important fields first
        priority_fields = ["name", "adresse", "steuernummer", "gesamtbetrag", "datum", "email", "telefon"]
        sorted_data = {}

        # Add priority fields first
        for field in priority_fields:
            if field in extracted_data:
                sorted_data[field] = extracted_data[field]

        # Add remaining fields
        for key, value in extracted_data.items():
            if key not in priority_fields:
                sorted_data[key] = value

        # Render data
        for key, value in sorted_data.items():
            if y_position < 50:
                p.showPage()
                p.setFont("Helvetica", 10)
                y_position = height - 50

            # Format key nicely
            key_formatted = key.replace("_", " ").title()

            # Handle long values
            value_str = str(value)
            if len(value_str) > 100:
                value_str = value_str[:100] + "..."

            p.setFont("Helvetica-Bold", 10)
            p.drawString(50, y_position, f"{key_formatted}:")
            p.setFont("Helvetica", 10)
            p.drawString(200, y_position, value_str)

            y_position -= 20

        # Footer
        p.line(50, 50, width - 50, 50)
        p.setFont("Helvetica-Oblique", 8)
        p.drawString(50, 35, "Dieses Dokument wurde automatisch erstellt. Bitte überprüfen Sie alle Angaben.")
        p.drawString(50, 25, "This document was automatically created. Please verify all information.")

        p.showPage()
        p.save()

        buffer.seek(0)

        logger.info(f"PDF created for {email} with {len(sorted_data)} fields")

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=steuerdokument_{email.split('@')[0]}.pdf"}
        )
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ADMIN TAB ENDPOINTS
@router.get("/admin/users")
async def get_all_users(db: Session = Depends(get_db), user: User = Depends(get_demo_user)):
    """Admin: Get all users with their cases"""
    # For MVP: Return demo data structure
    # In production: Query actual users

    cases = db.query(TaxCase).all()

    # Group by user email (for MVP, all belong to demo user)
    user_data = {
        "demo@example.com": {
            "email": "demo@example.com",
            "folders": [],
            "case_count": len(cases)
        }
    }

    return list(user_data.values())


# CASE MANAGEMENT ENDPOINTS
class CaseUpdateRequest(BaseModel):
    name: Optional[str] = None
    notes: Optional[str] = None

@router.put("/cases/{case_id}")
async def update_case(
    case_id: int,
    update_data: CaseUpdateRequest,
    db: Session = Depends(get_db),
    user: User = Depends(get_demo_user)
):
    """Update case (rename, edit notes)"""
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    if update_data.name:
        case.name = update_data.name
    if update_data.notes is not None:
        case.notes = update_data.notes

    db.commit()
    db.refresh(case)

    logger.info(f"Updated case {case_id}: name={update_data.name}")

    return {"success": True, "message": "Case updated"}


@router.delete("/cases/{case_id}")
async def delete_case(
    case_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_demo_user)
):
    """Delete a case and all associated data"""
    case = db.query(TaxCase).filter(
        TaxCase.id == case_id,
        TaxCase.user_id == user.id
    ).first()

    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    # Delete case (cascade will delete documents and extracted data)
    db.delete(case)
    db.commit()

    logger.info(f"Deleted case {case_id}")

    return {"success": True, "message": "Case deleted"}


@router.post("/free/export-pdf")
async def export_free_h7_pdf(
    extracted_data: Dict[str, Any],
    case_name: Optional[str] = None
):
    """
    Export extracted H7 data to Spanish PDF

    Args:
        extracted_data: Dictionary with extracted H7 field values
        case_name: Optional name for the document

    Returns:
        StreamingResponse with PDF file
    """
    try:
        from backend.services.document_service import generate_h7_pdf

        # Generate PDF
        pdf_buffer = generate_h7_pdf(extracted_data, case_name or "H7 Form")

        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"H7_Formulario_{timestamp}.pdf"

        # Return as streaming response
        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )

    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")
